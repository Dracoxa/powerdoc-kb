from __future__ import annotations

import base64
import io
import json
import math
import re
import tempfile
import time
import zipfile
from collections import Counter
from dataclasses import dataclass, field
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

import streamlit as st


APP_TITLE = "PowerDoc-KB"
SCHEMA_VERSION = "0.1.0"
ASSET_DIR = Path(__file__).resolve().parent / "assets"
ORBIT_LOOP_PATH = ASSET_DIR / "orbit-view-loop.webp"
ORBIT_VIDEO_PATH = ASSET_DIR / "orbit-background-2x-crossfade.mp4"

SYSTEM_SECTION_RULES = [
    ("mission_function", r"(任务和功能|任务需求|功能定义|功能分类|功能描述)"),
    ("orbit_environment", r"(轨道|环境|入轨高度|运行轨道|轨道周期|光照|阴影|地影|受晒因子)"),
    ("electrical_indicators", r"(电气技术指标|输出功率|母线电压|母线品质|火工品母线|抗扰|纹波|阶跃负载)"),
    ("topology_scheme", r"(拓扑方案|拓扑架构|光伏电源系统|功率通道|母线全调节|S3R|分流调节|放电调节)"),
    ("solar_array", r"(太阳电池翼|太阳电池阵|太阳翼|对日定向|展开|收拢|阵面面积|三结砷化镓)"),
    ("battery_energy_storage", r"(蓄电池|储能电池|锂离子|DOD|放电深度|循环寿命|深充放|恒流|恒压)"),
    ("control_equipment", r"(控制设备|电源管理器|分流调节器|充放电调节器|母线滤波器|驱动控制器|配电系统)"),
    ("reliability_maintenance", r"(可靠性|安全性|维修性|可测试性|在轨维修|寿命|故障|FMEA)"),
    ("mechanical_constraints", r"(重量|基频|共振|剩磁矩|磁干扰|外形尺寸|姿态控制|对接)"),
    ("document_deliverables", r"(研试文件|设计文件|设计报告|任务书|说明书|分析报告|文件名称)"),
]

SYSTEM_SECTION_LABELS = {
    "mission_function": "任务与功能",
    "orbit_environment": "轨道与环境边界",
    "electrical_indicators": "主要电气指标",
    "topology_scheme": "电源系统拓扑方案",
    "solar_array": "太阳电池翼与发电单元",
    "battery_energy_storage": "蓄电池与储能单元",
    "control_equipment": "控制设备与功率调节",
    "reliability_maintenance": "可靠性、安全性与在轨维修",
    "mechanical_constraints": "机械、磁与重量约束",
    "document_deliverables": "研试文件与设计交付物",
    "unclassified": "未归类条目",
}

TOPOLOGY_LABELS = {
    "buck_converter": "降压变换",
    "boost_converter": "升压变换",
    "buck_boost_converter": "升降压变换",
    "flyback_converter": "反激变换",
    "forward_converter": "正激变换",
    "half_bridge": "半桥",
    "full_bridge": "全桥",
    "llc_resonant": "LLC 谐振",
    "pfc": "功率因数校正",
    "ldo": "低压差稳压",
    "photovoltaic_power_system": "光伏电源系统",
    "regulated_primary_bus": "一次母线全调节",
    "s3r_shunt_regulator": "S3R 分流调节",
    "boost_discharge_regulator": "升压放电调节",
    "cc_cv_charging": "恒流转恒压充电",
    "multi_channel_power": "多功率通道",
    "dual_axis_solar_tracking": "双自由度对日定向",
}

RULE_CATEGORY_LABELS = {
    "layout": "布局与走线",
    "emi": "电磁兼容",
    "thermal": "热设计",
    "protection": "保护策略",
    "component_selection": "器件选型",
    "troubleshooting": "故障与调试",
    "system_topology": "系统拓扑",
    "bus_regulation": "母线调节",
    "power_capacity": "功率能力",
    "orbit_environment": "轨道与环境",
    "solar_array": "太阳电池翼",
    "energy_storage": "储能系统",
    "power_quality": "电能质量",
    "reliability_maintenance": "可靠性与维修",
    "mechanical_constraint": "结构与磁约束",
    "document_requirement": "交付文件",
}

TOPOLOGY_FEATURES = {
    "power_source": (r"(光伏电源系统|太阳电池翼|太阳电池阵|太阳翼)", "光伏电源系统 / 太阳电池翼发电"),
    "distribution_architecture": (r"(多功率通道|两个功率通道|功率通道相对独立|独立母线)", "多功率通道、通道相对独立"),
    "module_redundancy": (r"(多机组配置|四个相对独立的[“\"]?机组|四个机组)", "每通道多机组冗余"),
    "primary_bus": (r"(一次母线全调节|母线全调节|全调节方式|母线电压稳定)", "一次母线全调节"),
    "sunlight_control": (r"(S3R|顺序开关分流调节|分流调节)", "光照区 S3R 分流调节"),
    "eclipse_control": (r"(放电调节采用升压|升压控制调节|阴影区)", "阴影区储能电池升压放电调节"),
    "battery_charging": (r"(恒流转恒压|恒流|恒压充电)", "锂离子蓄电池恒流转恒压充电"),
    "solar_tracking": (r"(双自由度对日定向|对日定向系统|实时跟踪太阳)", "双自由度对日定向"),
}

TOPOLOGY_PATTERNS = {
    "buck_converter": r"\b(buck|step[-\s]?down|降压)\b",
    "boost_converter": r"\b(boost|step[-\s]?up|升压)\b",
    "buck_boost_converter": r"\b(buck[-\s]?boost|升降压)\b",
    "flyback_converter": r"\b(flyback|反激)\b",
    "forward_converter": r"\b(forward converter|正激)\b",
    "half_bridge": r"\b(half[-\s]?bridge|半桥)\b",
    "full_bridge": r"\b(full[-\s]?bridge|全桥)\b",
    "llc_resonant": r"\b(LLC|resonant converter|谐振)\b",
    "pfc": r"\b(PFC|power factor correction|功率因数校正)\b",
    "ldo": r"\b(LDO|low dropout|低压差)\b",
    "photovoltaic_power_system": r"(光伏电源系统|太阳电池翼|太阳电池阵|太阳翼)",
    "regulated_primary_bus": r"(一次母线全调节|母线全调节|全调节方式|母线电压稳定)",
    "s3r_shunt_regulator": r"(S3R|Sequential Switching Shunt Regulator|顺序开关分流调节|分流调节)",
    "boost_discharge_regulator": r"(放电调节采用升压|升压控制调节|放电调节)",
    "cc_cv_charging": r"(恒流转恒压|恒压充电|充电控制模式)",
    "multi_channel_power": r"(多功率通道|功率通道|多机组配置|独立母线)",
    "dual_axis_solar_tracking": r"(双自由度对日定向|对日定向系统|太阳翼实时跟踪)",
}

COMPONENT_PATTERNS = {
    "controller": r"\b(controller|PWM controller|控制器|控制芯片)\b",
    "mosfet": r"\b(MOSFET|FET|switching transistor|开关管)\b",
    "diode": r"\b(diode|Schottky|二极管|肖特基)\b",
    "inductor": r"\b(inductor|choke|电感)\b",
    "transformer": r"\b(transformer|变压器)\b",
    "input_capacitor": r"\b(input capacitor|Cin|输入电容)\b",
    "output_capacitor": r"\b(output capacitor|Cout|输出电容)\b",
    "current_sense": r"\b(current sense|sense resistor|电流检测|采样电阻)\b",
    "gate_driver": r"\b(gate driver|driver|栅极驱动|驱动器)\b",
    "snubber": r"\b(snubber|RC clamp|RCD clamp|吸收|钳位)\b",
    "solar_array": r"(太阳电池翼|太阳电池阵|太阳翼|太阳电池片)",
    "battery_pack": r"(蓄电池组|储能电池|锂离子蓄电池|电池组)",
    "shunt_regulator": r"(分流调节器|分流调节|S3R)",
    "charge_discharge_regulator": r"(充放电调节器|充电调节|放电调节)",
    "bus_filter": r"(母线滤波器|滤波器)",
    "power_management_unit": r"(电源管理器|配电系统|功率调配)",
    "solar_drive_controller": r"(驱动控制器|驱动机构|对日定向装置)",
    "pyro_bus": r"(火工品母线|火工品)",
}

RULE_CATEGORIES = {
    "layout": r"\b(layout|PCB|placement|place|route|trace|loop|ground|return path|布局|走线|接地|回路)\b",
    "emi": r"\b(EMI|EMC|noise|radiated|conducted|filter|shield|噪声|滤波|电磁)\b",
    "thermal": r"\b(thermal|temperature|heat|junction|derating|散热|温升|结温|降额)\b",
    "protection": r"\b(OCP|OVP|UVLO|short circuit|soft[-\s]?start|current limit|保护|限流|短路|软启动)\b",
    "component_selection": r"\b(select|selection|choose|rating|saturation|ESR|ripple current|选型|额定|饱和|纹波电流)\b",
    "troubleshooting": r"\b(failure|unstable|oscillation|ringing|overshoot|debug|troubleshoot|振荡|过冲|调试|故障)\b",
    "system_topology": r"(拓扑|架构|功率通道|机组|独立母线|配电系统|全调节|分流调节|放电调节)",
    "bus_regulation": r"(母线电压|一次母线|火工品母线|电压控制|稳定在|电压扰动|恢复时间)",
    "power_capacity": r"(输出功率|功耗|供电需求|发电功率|功率密度|比能量|能量平衡)",
    "orbit_environment": r"(轨道|入轨高度|运行轨道|轨道周期|光照|阴影|地影|受晒因子|太阳高度角)",
    "solar_array": r"(太阳电池翼|太阳电池阵|太阳翼|对日定向|展开|收拢|阵面面积|三结砷化镓)",
    "energy_storage": r"(蓄电池|储能电池|DOD|放电深度|循环寿命|深充放|恒流|恒压)",
    "power_quality": r"(母线品质|纹波|峰峰值|阶跃负载|抗扰|浪涌|跃变)",
    "reliability_maintenance": r"(可靠性|安全性|寿命|长寿命|在轨维修|故障|FMEA|可维修性|可测试性)",
    "mechanical_constraint": r"(基频|共振|姿态控制|对接|剩磁矩|磁干扰|外形尺寸|重量)",
    "document_requirement": r"(研试文件|设计文件|设计报告|任务书|技术说明书|使用说明书|分析报告)",
}

PARAM_ALIASES = {
    "input_voltage": ["vin", "v_in", "input voltage", "输入电压"],
    "output_voltage": ["vout", "v_out", "output voltage", "输出电压"],
    "output_current": ["iout", "i_out", "load current", "output current", "输出电流"],
    "switching_frequency": ["fsw", "f_sw", "switching frequency", "开关频率"],
    "efficiency": ["efficiency", "效率"],
    "output_ripple": ["ripple voltage", "output ripple", "纹波"],
    "inductance": ["inductance", "inductor value", "电感量"],
    "capacitance": ["capacitance", "capacitor value", "电容量"],
    "power": ["power", "output power", "功率", "输出功率", "功耗"],
    "bus_voltage": ["母线电压", "一次母线", "火工品母线"],
    "orbit_period": ["轨道周期", "周期"],
    "orbit_altitude": ["入轨高度", "运行轨道高度", "轨道高度"],
    "sunlight_duration": ["阳照区", "光照时间", "最短阳照区"],
    "eclipse_duration": ["阴影区", "最长阴影区"],
    "depth_of_discharge": ["DOD", "放电深度"],
    "ripple_voltage": ["纹波", "峰峰值", "p-p"],
    "recovery_time": ["恢复时间"],
    "surge_current": ["浪涌电流", "脉冲电流"],
    "cycle_life": ["循环寿命", "寿命"],
    "frequency": ["基频"],
    "area": ["阵面面积"],
    "efficiency": ["efficiency", "效率", "发电效率"],
}

PSEUDO_TABLE_LABELS = [
    "轨道倾角",
    "入轨高度",
    "运行轨道高度",
    "轨道周期",
    "光照/阴影时间",
    "受晒因子",
    "输出功率",
    "输出功率(单通道)",
    "母线电压",
    "火工品母线",
    "母线品质",
    "阶跃负载特性",
    "抗扰能力",
    "展开/收拢",
    "循环寿命",
    "发电效率",
    "磁特性",
    "对日定向",
    "深充放能力",
    "在轨更换恢复",
    "控制综合能力",
    "调节器独立能力",
    "分流与保护",
    "工作寿命",
    "可靠度指标",
    "总重上限",
    "功耗",
]

TOPOLOGY_CONTEXT_MARKERS = (
    "电源系统拓扑方案",
    "拓扑架构设计",
    "多机组冗余配置",
    "一次母线全调节控制逻辑",
    "一次母线采用全调节方式",
    "光照区控制机理",
    "光照区采用S3R",
    "阴影区控制机理",
    "阴影区储能电池通过升压放电调节",
    "区域平滑切换",
    "智能化充放电管理",
)

TOPOLOGY_NOISE_MARKERS = (
    "术语",
    "缩写",
    "BOL",
    "EOL",
    "DOD (",
    "GNC (",
    "S3R (",
    "PWR-TC-",
    "EQP-TC-",
    "考核",
    "验证",
    "测试人员操作",
)

UNIT_PATTERN = r"(V/ms|A/s|mV|V|A|W|kW|mW|Hz|kHz|MHz|uH|µH|mH|nF|uF|µF|mF|pF|mΩ|mohm|ohm|Ω|%|°C|C|km|m2|m²|s|ms|min|周次|次|年|Am²)"
VALUE_PATTERN = r"(?:[≥≤<>±]\s*)?[-+]?\d+(?:\.\d+)?(?:\s?(?:-|~|～|至|to)\s?(?:[≥≤<>±]\s*)?[-+]?\d+(?:\.\d+)?)?"


@dataclass
class ParsedDocument:
    name: str
    file_type: str
    text: str
    pages: list[str]
    backend: str = "legacy"
    structured_tables: list[dict[str, Any]] = field(default_factory=list)
    layout_items: list[dict[str, Any]] = field(default_factory=list)


def normalize_space(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def squeeze_ocr_spacing(text: str) -> str:
    text = normalize_space(text)
    text = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", "", text)
    text = re.sub(r"(?<=\d)\s+(?=\d)", "", text)
    text = re.sub(r"(?<=[A-Za-z])\s+(?=[A-Za-z])", "", text)
    text = re.sub(r"\s+([，。；：、）])", r"\1", text)
    text = re.sub(r"([（])\s+", r"\1", text)
    return text.strip()


def build_structured_table(table_item: Any) -> dict[str, Any]:
    data = getattr(table_item, "data", None)
    prov = getattr(table_item, "prov", []) or []
    page = prov[0].page_no if prov else None
    rows: list[list[str]] = []
    header: list[str] = []
    if data and getattr(data, "grid", None):
        for row_idx, row in enumerate(data.grid):
            row_text = [squeeze_ocr_spacing(getattr(cell, "text", "")) for cell in row]
            row_text = [cell for cell in row_text if cell]
            if not row_text:
                continue
            if row_idx == 0 and any(getattr(cell, "column_header", False) for cell in row):
                header = row_text
                continue
            rows.append(row_text)
    return {
        "page": page,
        "label": getattr(getattr(table_item, "label", None), "value", "table"),
        "header": header,
        "rows": rows,
    }


def extract_layout_items_from_docling(doc_obj: Any) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for payload in doc_obj.iterate_items():
        item = payload[0] if isinstance(payload, tuple) else payload
        label = getattr(getattr(item, "label", None), "value", type(item).__name__.lower())
        prov = getattr(item, "prov", []) or []
        page = prov[0].page_no if prov else None
        text = squeeze_ocr_spacing(getattr(item, "text", "") or getattr(item, "orig", "") or "")
        if label == "table":
            structured = build_structured_table(item)
            row_lines = [" | ".join(row) for row in structured["rows"] if row]
            header_line = " | ".join(structured["header"]) if structured["header"] else ""
            content = "\n".join(([header_line] if header_line else []) + row_lines)
            items.append(
                {
                    "kind": "table",
                    "page": structured["page"],
                    "label": label,
                    "text": content,
                    "table": structured,
                }
            )
            continue
        if not text:
            continue
        items.append(
            {
                "kind": "text",
                "page": page,
                "label": label,
                "text": text,
            }
        )
    return items


def read_docling_document(data: bytes, suffix: str) -> tuple[str, list[str], list[dict[str, Any]], list[dict[str, Any]]]:
    from docling.document_converter import DocumentConverter

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
        temp_file.write(data)
        temp_path = Path(temp_file.name)
    try:
        converter = DocumentConverter()
        result = converter.convert(str(temp_path))
        doc_obj = result.document
        layout_items = extract_layout_items_from_docling(doc_obj)
        structured_tables = [item["table"] for item in layout_items if item["kind"] == "table"]
        page_map: dict[int, list[str]] = {}
        for item in layout_items:
            page = item.get("page") or 1
            text = item.get("text", "")
            if not text:
                continue
            page_map.setdefault(page, []).append(text)
        if getattr(doc_obj, "pages", None):
            page_count = len(doc_obj.pages)
        else:
            page_count = max(page_map.keys(), default=1)
        pages = [normalize_space("\n\n".join(page_map.get(index, []))) for index in range(1, page_count + 1)]
        pages = [page for page in pages if page]
        full_text = normalize_space("\n\n".join(pages))
        return full_text, pages or [full_text], structured_tables, layout_items
    finally:
        temp_path.unlink(missing_ok=True)


def read_pdf(data: bytes) -> tuple[str, list[str]]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for page in reader.pages:
        pages.append(normalize_space(page.extract_text() or ""))
    return "\n\n".join(pages), pages


def read_docx(data: bytes) -> tuple[str, list[str]]:
    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
        chunks = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [normalize_space(cell.text.replace("\n", " ")) for cell in row.cells if cell.text.strip()]
                if cells:
                    chunks.append(" | ".join(cells))
        text = normalize_space("\n".join(chunks))
        return text, [text]
    except Exception:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            xml = archive.read("word/document.xml")
        root = ElementTree.fromstring(xml)
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        chunks = [node.text for node in root.findall(".//w:t", namespace) if node.text]
        text = normalize_space(" ".join(chunks))
        return text, [text]


def read_plain(data: bytes) -> tuple[str, list[str]]:
    for encoding in ("utf-8", "utf-16", "gb18030", "latin-1"):
        try:
            text = data.decode(encoding)
            return normalize_space(text), [normalize_space(text)]
        except UnicodeDecodeError:
            continue
    text = data.decode("utf-8", errors="ignore")
    return normalize_space(text), [normalize_space(text)]


def parse_upload(uploaded_file: Any) -> ParsedDocument:
    data = uploaded_file.getvalue()
    suffix = Path(uploaded_file.name).suffix.lower()
    backend = "legacy"
    structured_tables: list[dict[str, Any]] = []
    layout_items: list[dict[str, Any]] = []
    if suffix == ".pdf":
        file_type = "pdf"
        try:
            text, pages, structured_tables, layout_items = read_docling_document(data, suffix)
            backend = "docling"
        except Exception:
            text, pages = read_pdf(data)
    elif suffix == ".docx":
        file_type = "docx"
        try:
            text, pages, structured_tables, layout_items = read_docling_document(data, suffix)
            backend = "docling"
        except Exception:
            text, pages = read_docx(data)
    else:
        text, pages = read_plain(data)
        file_type = suffix.lstrip(".") or "text"
    return ParsedDocument(
        name=uploaded_file.name,
        file_type=file_type,
        text=normalize_space(text),
        pages=[normalize_space(page) for page in pages],
        backend=backend,
        structured_tables=structured_tables,
        layout_items=layout_items,
    )


def split_sentences(text: str) -> list[str]:
    raw = re.split(r"(?<=[。！？.!?])\s+|\n+", text)
    return [item.strip(" -•\t") for item in raw if len(item.strip()) > 12]


def split_table_rows(text: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in text.splitlines():
        if "|" not in line:
            continue
        cells = [normalize_space(cell) for cell in line.split("|")]
        cells = [cell for cell in cells if cell and cell not in {"-", "—"}]
        if len(cells) >= 2:
            rows.append(cells)
    return rows


def is_header_row(cells: list[str]) -> bool:
    joined = "".join(cells)
    header_tokens = {
        "功能分类功能描述关键配置/要求",
        "参数项指标数值/描述备注",
        "指标类别具体参数要求工况/约束条件",
        "特性维度指标要求设计意义",
        "序号类别文件名称文件类型",
        "序号类别项目类别备注",
        "项目指标详情",
    }
    return joined in header_tokens or all(not re.search(r"\d|≥|≤|V|W|Hz|DOD|FMEA", cell, flags=re.IGNORECASE) for cell in cells[1:])


def document_blocks(text: str) -> list[str]:
    blocks = []
    for line in text.splitlines():
        item = normalize_space(line)
        if len(item) >= 2:
            blocks.append(item)
    if blocks:
        return blocks
    return split_sentences(text)


def document_blocks_from_layout(doc: ParsedDocument) -> list[str]:
    if not doc.layout_items:
        return document_blocks(doc.text)
    blocks: list[str] = []
    for item in doc.layout_items:
        if item.get("kind") != "text":
            continue
        text = squeeze_ocr_spacing(item.get("text", ""))
        if len(text) >= 2:
            blocks.append(text)
    return blocks or document_blocks(doc.text)


def classify_system_section(text: str) -> str:
    for section, pattern in SYSTEM_SECTION_RULES:
        if re.search(pattern, text, flags=re.IGNORECASE):
            return section
    return "unclassified"


def parameter_name_from_label(label: str) -> str:
    label_norm = label.lower()
    for canonical, aliases in PARAM_ALIASES.items():
        if any(alias.lower() in label_norm for alias in aliases):
            return canonical
    cleaned = re.sub(r"[^\w\u4e00-\u9fff]+", "_", label.strip().lower()).strip("_")
    return cleaned[:48] or "parameter"


def source_for_text(pages: list[str], needle: str) -> dict[str, Any]:
    needle_norm = normalize_space(needle)
    for index, page in enumerate(pages, start=1):
        if needle_norm[:80] and needle_norm[:80] in normalize_space(page):
            return {"page": index, "snippet": needle_norm[:420]}
    return {"page": None, "snippet": needle_norm[:420]}


def confidence_from_hits(*hits: bool) -> float:
    base = 0.52 + sum(0.11 for hit in hits if hit)
    return round(min(base, 0.93), 2)


def clean_heading(text: str) -> str:
    text = normalize_space(text)
    text = re.sub(r"^[一二三四五六七八九十]+[、.．]\s*", "", text)
    text = re.sub(r"^\d+[、.．]\s*", "", text)
    return text


def compact_text(text: str, limit: int = 96) -> str:
    text = normalize_space(text).strip("，；。 ")
    if len(text) <= limit:
        return text
    clipped = text[:limit]
    for marker in ("。", "；", "，"):
        if marker in clipped[-24:]:
            clipped = clipped.rsplit(marker, 1)[0]
            break
    return clipped.rstrip("，；。 ") + "…"


def build_source(page: int | None, snippet: str) -> dict[str, Any]:
    return {
        "page": page,
        "snippet": compact_text(snippet, 64),
    }


def make_source(doc: ParsedDocument, text: str) -> dict[str, Any]:
    return source_for_text(doc.pages, text)


def topology_context_pages(doc: ParsedDocument) -> list[dict[str, Any]]:
    contexts: list[dict[str, Any]] = []
    for index, page in enumerate(doc.pages, start=1):
        page_norm = normalize_space(page)
        if any(marker in page_norm for marker in TOPOLOGY_CONTEXT_MARKERS):
            contexts.append({"page": index, "text": page_norm})
    if contexts:
        return contexts
    sections = extract_system_sections(doc)
    fallback: list[dict[str, Any]] = []
    for item in sections.get("topology_scheme", []):
        snippet = normalize_space(item["content"])
        if any(marker in snippet for marker in TOPOLOGY_NOISE_MARKERS):
            continue
        fallback.append({"page": item["source"].get("page"), "text": snippet})
    return fallback


def clean_topology_snippet(text: str) -> str:
    cleaned = normalize_space(text)
    cleaned = re.sub(r"[A-Z]{2,}\s*\([^)]{2,}\)[：:][^。；]*", "", cleaned)
    cleaned = re.sub(r"\b(?:PWR|EQP)-TC-\d+\b[^。；]*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return compact_text(cleaned, 88)


def summarize_topology_feature(feature_id: str, text: str) -> str:
    snippet = clean_topology_snippet(text)
    phrase_patterns = {
        "power_source": r"(以单侧太阳电池翼为能量起点|太阳电池翼作为核心发电设备|太阳电池阵被划分为多个独立的太阳电池阵段)",
        "distribution_architecture": r"(两个完全对称、相互独立的功率通道|高度模块化的多功率通道光伏电源系统架构|通道内其余三个机组仍能提供至少75%的额定功率)",
        "module_redundancy": r"(每个通道配置四个相对独立的[“\"]?机组[”\"]?|四个机组在物理上并联输出至通道一次母线|无集中单点故障)",
        "primary_bus": r"(系统采用[“\"]?一次母线全调节拓扑[”\"]?|实现一次母线全调节控制|母线电压被牢牢控制在103V~105V的高端区)",
        "sunlight_control": r"(光照区[^。；]*S3R[^。；]*|采用S3R型功率调节技术[^。；]*)",
        "eclipse_control": r"(阴影区[^。；]*升压[^。；]*|放电调节采用高效率的同步升压[^。；]*)",
        "battery_charging": r"(恒流转恒压[^。；]*控制模式|系统自动无缝转入高精度恒压充电模式)",
        "solar_tracking": r"(双自由度对日定向|控制太阳电池翼双自由度转动)",
    }
    pattern = phrase_patterns.get(feature_id)
    if pattern:
        match = re.search(pattern, snippet)
        if match:
            return compact_text(match.group(1), 56)
    fallback_labels = {
        "power_source": "主能源来自太阳电池翼发电链路。",
        "distribution_architecture": "系统采用相互独立的多功率通道架构。",
        "module_redundancy": "每个功率通道内部继续做多机组冗余。",
        "primary_bus": "一次母线采用全调节控制。",
        "sunlight_control": "光照区由 S3R 分流调节维持母线稳定。",
        "eclipse_control": "阴影区由升压放电调节维持母线稳定。",
        "battery_charging": "蓄电池采用恒流转恒压充电策略。",
        "solar_tracking": "太阳翼具备双自由度对日定向能力。",
    }
    return fallback_labels.get(feature_id, snippet)


def contains_quantity(text: str) -> bool:
    return bool(re.search(rf"{VALUE_PATTERN}\s*(?:{UNIT_PATTERN})?", text, flags=re.IGNORECASE))


def parse_value_ranges(text: str) -> list[dict[str, Any]]:
    ranges = []
    cleaned_text = text.replace(",", "")
    range_pattern = re.compile(
        rf"(?P<lower>[≥≤<>±]?\s*[-+]?\d+(?:\.\d+)?)\s*(?P<unit1>{UNIT_PATTERN})?\s*(?:-|~|～|至|to)\s*(?P<upper>[≥≤<>±]?\s*[-+]?\d+(?:\.\d+)?)\s*(?P<unit2>{UNIT_PATTERN})?",
        flags=re.IGNORECASE,
    )
    scalar_pattern = re.compile(
        rf"(?P<value>[≥≤<>±]?\s*[-+]?\d+(?:\.\d+)?)\s*(?P<unit>{UNIT_PATTERN})",
        flags=re.IGNORECASE,
    )
    for match in range_pattern.finditer(cleaned_text):
        unit = match.group("unit2") or match.group("unit1") or ""
        ranges.append(
            {
                "kind": "range",
                "lower": re.sub(r"\s+", "", match.group("lower")),
                "upper": re.sub(r"\s+", "", match.group("upper")),
                "unit": unit,
                "text": normalize_space(match.group(0)),
            }
        )
    if ranges:
        return ranges
    for match in scalar_pattern.finditer(cleaned_text):
        ranges.append(
            {
                "kind": "scalar",
                "value": re.sub(r"\s+", "", match.group("value")),
                "unit": match.group("unit"),
                "text": normalize_space(match.group(0)),
            }
        )
    return ranges


def normalize_cells(line: str) -> list[str]:
    cells = [normalize_space(cell) for cell in line.split("|")]
    return [cell for cell in cells if cell and cell not in {"-", "—"}]


def row_text_from_cells(cells: list[str]) -> str:
    return " | ".join(squeeze_ocr_spacing(cell) for cell in cells if cell)


def extract_table_blocks(doc: ParsedDocument) -> list[dict[str, Any]]:
    if doc.structured_tables:
        blocks: list[dict[str, Any]] = []
        for table in doc.structured_tables:
            header = [squeeze_ocr_spacing(cell) for cell in table.get("header", []) if cell]
            rows = []
            for row in table.get("rows", []):
                cleaned_row = [squeeze_ocr_spacing(cell) for cell in row if cell]
                if len(cleaned_row) >= 2:
                    rows.append(cleaned_row)
            if not rows:
                continue
            section_seed = " | ".join(header) if header else " | ".join(rows[0])
            blocks.append(
                {
                    "heading": "",
                    "header": header,
                    "rows": rows,
                    "source": build_source(table.get("page"), row_text_from_cells(rows[0])),
                    "section": classify_system_section(section_seed),
                }
            )
        if blocks:
            return blocks
    blocks: list[dict[str, Any]] = []
    current: list[str] = []
    for raw_line in doc.text.splitlines():
        line = normalize_space(raw_line)
        if "|" in line:
            current.append(line)
            continue
        if current:
            blocks.append(build_table_block(doc, current))
            current = []
    if current:
        blocks.append(build_table_block(doc, current))
    return [block for block in blocks if block["rows"]]


def build_table_block(doc: ParsedDocument, lines: list[str]) -> dict[str, Any]:
    rows = [normalize_cells(line) for line in lines if "|" in line]
    rows = [row for row in rows if row]
    header: list[str] = []
    body = rows[:]
    if rows and is_header_row(rows[0]):
        header = rows[0]
        body = rows[1:]
    elif rows and len(rows[0]) >= 2 and all(not contains_quantity(cell) for cell in rows[0]):
        header = rows[0]
        body = rows[1:]
    heading = ""
    first_line = normalize_space(lines[0]) if lines else ""
    if "|" not in first_line:
        heading = first_line
    section_seed = " | ".join(header) if header else " | ".join(rows[0]) if rows else ""
    return {
        "heading": heading,
        "header": header,
        "rows": body,
        "source": make_source(doc, "\n".join(lines)),
        "section": classify_system_section(section_seed),
    }


def split_pseudo_table_segments(text: str) -> list[str]:
    pattern = re.compile(
        "(" + "|".join(sorted((re.escape(label) for label in PSEUDO_TABLE_LABELS), key=len, reverse=True)) + r")"
    )
    matches = list(pattern.finditer(text))
    if not matches:
        return []
    segments = []
    for idx, match in enumerate(matches):
        start = match.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        segment = normalize_space(text[start:end])
        if len(segment) >= 6:
            segments.append(segment)
    return segments


def pseudo_candidate_from_segment(doc: ParsedDocument, segment: str) -> dict[str, Any] | None:
    if len(segment) > 260 and not segment.startswith(tuple(PSEUDO_TABLE_LABELS)):
        return None
    label_match = re.match(
        "(" + "|".join(sorted((re.escape(label) for label in PSEUDO_TABLE_LABELS), key=len, reverse=True)) + r")",
        segment,
    )
    if not label_match:
        return None
    label = label_match.group(1)
    remainder = normalize_space(segment[len(label):].lstrip("：: "))
    if not remainder or not contains_quantity(remainder):
        return None
    stop_markers = [
        "设计意义在于",
        "工况与约束条件",
        "设计边界与备注",
        "主要考验",
        "该项作为",
        "以满足",
        "以防",
        "验证",
        "系统处于",
        "模拟",
        "控制与核心调节设备",
        "B.",
        "C.",
        "4.",
        "5.",
        "三、",
    ]
    for marker in stop_markers:
        if marker in remainder:
            remainder = remainder.split(marker, 1)[0].strip("；，。 ")
    if label == "输出功率" and "储能电池放电深度" in remainder:
        remainder = remainder.split("储能电池放电深度", 1)[0].strip("；，。 ")
    if label == "发电效率" and "剩磁矩" in remainder:
        remainder = remainder.split("剩磁矩", 1)[0].strip("；，。 ")
    if label == "母线品质" and "测量带宽" in remainder and "纹波" in remainder:
        remainder = remainder.split("测量带宽", 1)[0].strip("；，。 ")
    if label == "抗扰能力" and "一次母线电压扰动" in remainder:
        remainder = remainder.split("一次母线电压扰动", 1)[0] + " 一次母线电压扰动"
    quantities = parse_value_ranges(remainder)
    if not quantities:
        return None
    if len(remainder) > 220 and "。" in remainder:
        remainder = remainder.split("。", 1)[0]
    return {
        "label": clean_heading(label),
        "value": remainder,
        "cells": [label, remainder],
        "row_text": f"{label} | {remainder}",
        "header": [],
        "section": classify_system_section(f"{label} {remainder}"),
        "source": make_source(doc, segment),
        "quantities": quantities,
    }


def extract_pseudo_table_candidates(doc: ParsedDocument) -> list[dict[str, Any]]:
    if doc.structured_tables:
        return []
    candidates: list[dict[str, Any]] = []
    seen: set[str] = set()
    for page in doc.pages:
        for segment in split_pseudo_table_segments(page):
            candidate = pseudo_candidate_from_segment(doc, segment)
            if not candidate:
                continue
            key = f"{candidate['label']}::{candidate['value'][:120]}"
            if key in seen:
                continue
            seen.add(key)
            candidates.append(candidate)
    return candidates


def extract_table_candidates(doc: ParsedDocument) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    seen: set[str] = set()
    for block in extract_table_blocks(doc):
        for row in block["rows"]:
            if not row:
                continue
            label = clean_heading(row[0])
            note = ""
            if doc.structured_tables and len(row) >= 3:
                value = row[1].strip()
                note = "；".join(row[2:]).strip()
            else:
                value = "；".join(row[1:]).strip() if len(row) > 1 else ""
            row_text = " | ".join(row)
            candidates.append(
                {
                    "label": label,
                    "value": value,
                    "note": note,
                    "cells": row,
                    "row_text": row_text,
                    "header": block["header"],
                    "section": classify_system_section(f"{label} | {value} | {note} | {' | '.join(block['header'])}"),
                    "source": make_source(doc, row_text),
                    "quantities": parse_value_ranges(value),
                }
            )
    merged = []
    for candidate in candidates + extract_pseudo_table_candidates(doc):
        key = f"{candidate['label']}::{candidate['value'][:140]}"
        if key in seen:
            continue
        seen.add(key)
        label = candidate["label"]
        if label in {"输出功率", "输出功率(单通道)", "输出功率\n(单通道)", "母线电压", "火工品母线", "母线品质", "阶跃负载特性", "抗扰能力"}:
            candidate["section"] = "electrical_indicators"
        elif label in {"循环寿命", "深充放能力"}:
            candidate["section"] = "battery_energy_storage"
        elif label in {"发电效率", "磁特性", "展开/收拢", "对日定向"}:
            candidate["section"] = "solar_array"
        elif label in {"工作寿命", "可靠度指标"}:
            candidate["section"] = "reliability_maintenance"
        elif label in {"在轨更换、恢复能力"}:
            candidate["section"] = "battery_energy_storage"
        elif label in {"电源分系统重量", "自身功耗"}:
            candidate["section"] = "mechanical_constraints"
        elif label in {"可靠度", "系统寿命", "地面总装测试时间"}:
            candidate["section"] = "reliability_maintenance"
        elif label in {"充放电调节器", "分流调节器", "过压保护", "控制调节及管理设备要求", "故障隔离、在轨管理"}:
            candidate["section"] = "control_equipment"
        merged.append(candidate)
    return merged


def is_deliverable_candidate(candidate: dict[str, Any]) -> bool:
    row_text = candidate["row_text"]
    header_text = " ".join(candidate.get("header", []))
    return (
        "文件名称" in header_text
        or "文件类型" in header_text
        or candidate["section"] == "document_deliverables"
        or bool(re.search(RULE_CATEGORIES["document_requirement"], row_text))
    )


def is_test_case_text(text: str) -> bool:
    return bool(re.search(r"\b(?:PWR|EQP)-TC-\d+\b", text, flags=re.IGNORECASE))


def normalize_test_case_id(text: str) -> str:
    compact = squeeze_ocr_spacing(text).replace(" ", "")
    match = re.search(r"(?i)(PWR|EQP)-?TC-?(\d+)", compact)
    if not match:
        return compact
    return f"{match.group(1).upper()}-TC-{int(match.group(2)):02d}"


def extract_test_cases(doc: ParsedDocument) -> list[dict[str, Any]]:
    test_cases: list[dict[str, Any]] = []
    seen: set[str] = set()
    continuation: dict[str, Any] | None = None

    def append_fragments(case: dict[str, Any], row: list[str]) -> None:
        cleaned = [squeeze_ocr_spacing(cell) for cell in row]
        if len(cleaned) >= 6:
            if cleaned[1]:
                case["name"] = f"{case['name']} {cleaned[1]}".strip()
            if cleaned[2]:
                case["precondition"] = f"{case['precondition']} {cleaned[2]}".strip()
            if cleaned[3]:
                case["steps"] = f"{case['steps']} {cleaned[3]}".strip()
            if cleaned[4]:
                case["target_rule"] = f"{case['target_rule']} {cleaned[4]}".strip()
            if cleaned[5]:
                case["acceptance_criteria"] = f"{case['acceptance_criteria']} {cleaned[5]}".strip()
            return
        if len(cleaned) == 5:
            if cleaned[0]:
                case["name"] = f"{case['name']} {cleaned[0]}".strip()
            if cleaned[1]:
                case["precondition"] = f"{case['precondition']} {cleaned[1]}".strip()
            if cleaned[2]:
                case["steps"] = f"{case['steps']} {cleaned[2]}".strip()
            if cleaned[3]:
                case["target_rule"] = f"{case['target_rule']} {cleaned[3]}".strip()
            if cleaned[4]:
                case["acceptance_criteria"] = f"{case['acceptance_criteria']} {cleaned[4]}".strip()
            return
        if cleaned and cleaned[-1]:
            case["acceptance_criteria"] = f"{case['acceptance_criteria']} {cleaned[-1]}".strip()

    for table in doc.structured_tables:
        header_text = " ".join(table.get("header", []))
        if "项目编号" not in header_text and "测试项名称" not in header_text:
            if continuation:
                for row in table.get("rows", []):
                    if row and re.match(r"^(?:PWR|EQP)-TC-\d+$", normalize_test_case_id(row[0])):
                        continuation = None
                        break
                    append_fragments(continuation, row)
            continue
        for row in table.get("rows", []):
            normalized_id = normalize_test_case_id(row[0]) if row else ""
            if continuation and not re.match(r"^(?:PWR|EQP)-TC-\d+$", normalized_id):
                append_fragments(continuation, row)
                continue
            if len(row) < 6:
                continue
            case_id = normalized_id
            if not re.match(r"^(?:PWR|EQP)-TC-\d+$", case_id):
                continue
            key = f"{table.get('page')}::{case_id}"
            if key in seen:
                continue
            seen.add(key)
            case = {
                "id": case_id,
                "name": squeeze_ocr_spacing(row[1]),
                "precondition": squeeze_ocr_spacing(row[2]),
                "steps": squeeze_ocr_spacing(row[3]),
                "target_rule": squeeze_ocr_spacing(row[4]),
                "acceptance_criteria": squeeze_ocr_spacing(row[5]),
                "source": build_source(table.get("page"), " | ".join(row)),
            }
            test_cases.append(case)
            continuation = case
    return test_cases[:80]


def is_rule_sentence(sentence: str) -> bool:
    if len(sentence) < 18:
        return False
    if re.match(r"^[一二三四五六七八九十0-9]+[、.．]?\s*[^，。]{0,24}$", sentence):
        return False
    if re.search(r"(术语|缩写|定义说明|引言与概述|适用范围)", sentence):
        return False
    if re.search(r"(BOL|EOL|DOD|GNC|S3R)\s*\(", sentence) and "必须" not in sentence and "要求" not in sentence:
        return False
    return True


def format_quantities(quantities: list[dict[str, Any]], raw_text: str) -> tuple[str, list[str]]:
    if not quantities:
        units = sorted(set(re.findall(UNIT_PATTERN, raw_text, flags=re.IGNORECASE)))
        return raw_text, units[:6]
    display_parts = []
    units = []
    for item in quantities:
        unit = item.get("unit") or ""
        if unit:
            units.append(unit)
        if item["kind"] == "range":
            display_parts.append(f"{item['lower']}~{item['upper']}{unit}")
        else:
            display_parts.append(f"{item['value']}{unit}")
    return "；".join(display_parts), sorted(set(units))[:6]


def split_metric_details(label: str, value: str, note: str = "") -> list[dict[str, Any]]:
    value = squeeze_ocr_spacing(value)
    note = squeeze_ocr_spacing(note)
    label = squeeze_ocr_spacing(label)
    details: list[dict[str, Any]] = []

    def push(name: str, text: str, detail_note: str = "") -> None:
        text = squeeze_ocr_spacing(text)
        if not text:
            return
        quantities = parse_value_ranges(text)
        numeric_hint = bool(re.search(r"\d", text))
        if "Am²" in text or "m²" in text or "℃" in text:
            value_display = re.sub(r"\s+", "", text)
            units = []
            if "Am²" in text:
                units = ["Am²"]
            elif "m²" in text:
                units = ["m²"]
            elif "℃" in text:
                units = ["℃"]
        elif re.search(r"\d+\s*±\s*\d+", text):
            value_display = re.sub(r"\s+", "", text)
            units = sorted(set(re.findall(UNIT_PATTERN, text, flags=re.IGNORECASE)))[:6]
        else:
            value_display, units = format_quantities(quantities, text)
        details.append(
            {
                "name": name,
                "value": value_display,
                "raw_value": text,
                "note": detail_note or note,
                "units": units,
                "has_numeric_value": bool(quantities) or numeric_hint,
            }
        )

    if label == "入轨高度":
        near_match = re.search(r"近地点\s*([^\s，；]+(?:\s*km)?)", value, flags=re.IGNORECASE)
        far_match = re.search(r"远地点\s*([^\s，；]+(?:\s*km)?)", value, flags=re.IGNORECASE)
        if near_match or far_match:
            if near_match:
                push("入轨高度/近地点", near_match.group(1))
            if far_match:
                push("入轨高度/远地点", far_match.group(1))
            return details

    if label == "光照/阴影时间":
        sunlight_match = re.search(r"最短阳照区[:：]?\s*([^\s，；]+(?:\s*min)?)", value, flags=re.IGNORECASE)
        eclipse_match = re.search(r"最长阴影区[:：]?\s*([^\s，；]+(?:\s*min)?)", value, flags=re.IGNORECASE)
        if sunlight_match or eclipse_match:
            if sunlight_match:
                push("光照时间/最短阳照区", sunlight_match.group(1))
            if eclipse_match:
                push("阴影时间/最长阴影区", eclipse_match.group(1))
            return details

    if label in {"输出功率(单通道)", "输出功率\n(单通道)"}:
        named_patterns = [
            ("输出功率(单通道)/一次展开", r"≥\s*1\.?8\s*kW"),
            ("输出功率(单通道)/BOL", r"≥\s*7\.?0\s*kW"),
            ("输出功率(单通道)/EOL", r"≥\s*6\.?2\s*kW"),
            ("输出功率(单通道)/稳态放电", r"≥\s*11\s*kW"),
        ]
        found = False
        for name, pattern in named_patterns:
            match = re.search(pattern, value, flags=re.IGNORECASE)
            if match:
                push(name, match.group(0))
                found = True
        if found:
            return details

    if label == "运行轨道高度":
        center_range = re.search(r"(\d+\s*±\s*\d+\s*km)", value, flags=re.IGNORECASE)
        if center_range:
            push(label, center_range.group(1))
            return details

    if label == "轨道周期":
        main_match = re.search(r"(\d+\s*s)", value, flags=re.IGNORECASE)
        approx_match = re.search(r"(约\s*\d+(?:\.\d+)?\s*分钟)", value, flags=re.IGNORECASE)
        if main_match:
            push("轨道周期/秒", main_match.group(1))
        if approx_match:
            push("轨道周期/分钟", approx_match.group(1))
        if details:
            return details

    if label.startswith("母线品质"):
        ripple_match = re.search(r"(≤\s*[\d.]+\s*mV)", value, flags=re.IGNORECASE)
        recovery_match = re.search(r"(恢复时间\s*≤\s*[\d.]+\s*ms)", value, flags=re.IGNORECASE)
        bandwidth_match = re.search(r"(0\s*[~～]\s*10\s*MHz)", note, flags=re.IGNORECASE)
        if ripple_match:
            push(f"{label}/纹波上限", ripple_match.group(1), note)
        if recovery_match:
            push(f"{label}/恢复时间", recovery_match.group(1), note)
        if bandwidth_match:
            push(f"{label}/测量带宽", bandwidth_match.group(1), note)
        if details:
            return details

    if label == "受晒因子":
        pairs = re.findall(r"(\d+)\s*°\s*:\s*([\d.]+)", value)
        if pairs:
            for angle, factor in pairs:
                push(f"受晒因子/{angle}°", factor)
            return details

    if label == "循环寿命":
        cycle_match = re.search(r"(≥\s*[\d,]+\s*周次)", value, flags=re.IGNORECASE)
        dod_match = re.search(r"(\d+\s*%\s*DOD)", value, flags=re.IGNORECASE)
        if cycle_match:
            push("循环寿命/周次", cycle_match.group(1))
        if dod_match:
            push("循环寿命/DOD", dod_match.group(1))
        if details:
            return details

    if label == "深充放能力":
        dod_match = re.search(r"(DOD\s*≤\s*\d+\s*%)", value, flags=re.IGNORECASE)
        window_match = re.search(r"(\d+\s*年累计)", value, flags=re.IGNORECASE)
        count_match = re.search(r"(≤\s*\d+\s*次)", value, flags=re.IGNORECASE)
        if dod_match:
            push("深充放能力/DOD上限", dod_match.group(1))
        if window_match:
            push("深充放能力/统计窗口", window_match.group(1))
        if count_match:
            push("深充放能力/累计次数上限", count_match.group(1))
        if details:
            return details

    if label == "在轨更换、恢复能力":
        replace_match = re.search(r"(\d+\s*年内可更换)", value, flags=re.IGNORECASE)
        temp_match = re.search(r"(\d+\s*[~～]\s*\d+\s*℃)", value)
        humidity_match = re.search(r"(\d+\s*%\s*[~～]\s*\d+\s*%)", value)
        if replace_match:
            push("在轨更换/时限", replace_match.group(1))
        if temp_match:
            push("在轨更换/存储温度", temp_match.group(1))
        if humidity_match:
            push("在轨更换/存储湿度", humidity_match.group(1))
        if details:
            return details

    if label == "自身功耗":
        solo_match = re.search(r"(单飞行[:：]?\s*≤\s*[\d.]+\s*W)", value, flags=re.IGNORECASE)
        docked_match = re.search(r"(组合体飞行[:：]?\s*≤\s*[\d.]+\s*W)", value, flags=re.IGNORECASE)
        if solo_match:
            push("自身功耗/单飞行", solo_match.group(1))
        if docked_match:
            push("自身功耗/组合体飞行", docked_match.group(1))
        if details:
            return details

    if label == "可靠度":
        normalized_reliability = value.replace("年全寿命期", " 13年全寿命期")
        meet_match = re.search(r"(交会对接期\s*\(4天\)[:：]?\s*≥\s*[\d.]+)", normalized_reliability, flags=re.IGNORECASE)
        full_life_match = re.search(r"(13\s*年全寿命期[:：]?\s*≥\s*[\d.]+)", normalized_reliability, flags=re.IGNORECASE)
        if meet_match:
            push("可靠度/交会对接期", meet_match.group(1))
        if full_life_match:
            push("可靠度/13年全寿命期", full_life_match.group(1))
        if details:
            return details

    if label == "磁特性":
        magnetic_match = re.search(r"(≤\s*[\d.]+\s*Am²)", value, flags=re.IGNORECASE)
        if magnetic_match:
            push("磁特性/剩磁矩", magnetic_match.group(1), note)
            return details

    if label == "母线品质":
        ripple_match = re.search(r"(纹波\s*≤\s*[\d.]+\s*mV)", value, flags=re.IGNORECASE)
        recovery_match = re.search(r"(恢复时间\s*≤\s*[\d.]+\s*ms)", value, flags=re.IGNORECASE)
        if ripple_match or recovery_match:
            if ripple_match:
                push("母线品质/纹波", ripple_match.group(1))
            if recovery_match:
                push("母线品质/恢复时间", recovery_match.group(1))
            return details

    if label == "火工品母线":
        voltage_match = re.search(r"(\d+\s*V\s*(?:-|~|～|至)\s*\d+\s*V)", value, flags=re.IGNORECASE)
        pulse_match = re.search(r"(?:支持)?\s*(0\s*A\s*~\s*60\s*A)", note or value, flags=re.IGNORECASE)
        if voltage_match:
            push("火工品母线/电压范围", voltage_match.group(1))
        if pulse_match:
            push("火工品母线/脉冲电流", pulse_match.group(1), note)
        if details:
            return details

    push(label, value, note)
    return details


def row_to_metric(cells: list[str], doc: ParsedDocument) -> dict[str, Any] | None:
    if is_header_row(cells):
        return None
    label = clean_heading(cells[0])
    value = "；".join(cells[1:]).strip()
    if not label or not value:
        return None
    section = classify_system_section(" | ".join(cells))
    units = sorted(set(re.findall(UNIT_PATTERN, value, flags=re.IGNORECASE)))
    numeric_values = re.findall(rf"{VALUE_PATTERN}\s*(?:{UNIT_PATTERN})?", value, flags=re.IGNORECASE)
    return {
        "name": label,
        "section": section,
        "value": value,
        "units": units[:6],
        "has_numeric_value": bool(numeric_values),
        "source": make_source(doc, " | ".join(cells)),
    }


def extract_system_sections(doc: ParsedDocument) -> dict[str, list[dict[str, Any]]]:
    sections: dict[str, list[dict[str, Any]]] = {key: [] for key, _ in SYSTEM_SECTION_RULES}
    sections["unclassified"] = []
    seen: set[str] = set()
    for candidate in extract_table_candidates(doc):
        row_text = candidate["row_text"]
        section = candidate["section"]
        key = re.sub(r"\W+", "", row_text.lower())[:140]
        if key in seen:
            continue
        seen.add(key)
        sections.setdefault(section, []).append(
            {
                "kind": "table_row",
                "title": candidate["label"],
                "content": candidate["value"] or row_text,
                "cells": candidate["cells"],
                "source": candidate["source"],
            }
        )
    for block in document_blocks_from_layout(doc):
        if "|" in block or len(block) < 18:
            continue
        section = classify_system_section(block)
        if section == "unclassified":
            continue
        key = re.sub(r"\W+", "", block.lower())[:140]
        if key in seen:
            continue
        seen.add(key)
        sections.setdefault(section, []).append(
            {
                "kind": "paragraph",
                "title": clean_heading(block[:34]),
                "content": block,
                "cells": [],
                "source": make_source(doc, block),
            }
        )
    return {section: items for section, items in sections.items() if items}


def extract_structured_metrics(doc: ParsedDocument) -> list[dict[str, Any]]:
    metrics: list[dict[str, Any]] = []
    seen: set[str] = set()
    for candidate in extract_table_candidates(doc):
        if is_deliverable_candidate(candidate):
            continue
        if is_test_case_text(candidate["value"]) or is_test_case_text(candidate["row_text"]):
            continue
        if not candidate["quantities"] and candidate["section"] not in {"document_deliverables", "mission_function"}:
            continue
        note = candidate.get("note", "")
        metric_parts = split_metric_details(candidate["label"], candidate["value"], note)
        for part in metric_parts:
            if not part["value"]:
                continue
            metric = {
                "name": part["name"],
                "section": candidate["section"],
                "value": part["value"],
                "raw_value": part["raw_value"],
                "note": part["note"],
                "units": part["units"],
                "has_numeric_value": part["has_numeric_value"],
                "source": candidate["source"],
            }
            key = f"{metric['section']}::{metric['name']}::{metric['value']}::{metric.get('note','')}"
            if key in seen:
                continue
            seen.add(key)
            metrics.append(metric)
    return metrics[:160]


def extract_functional_requirements(doc: ParsedDocument) -> list[dict[str, Any]]:
    functions: list[dict[str, Any]] = []
    seen: set[str] = set()
    for candidate in extract_table_candidates(doc):
        if candidate["section"] != "mission_function":
            continue
        label = candidate["label"]
        if label in {"在轨拓展要求"} and candidate["value"].strip("—- " or "") == "":
            continue
        key = f"{label}::{candidate['value']}::{candidate.get('note','')}"
        if key in seen:
            continue
        seen.add(key)
        functions.append(
            {
                "name": label,
                "description": squeeze_ocr_spacing(candidate["value"]),
                "configuration": squeeze_ocr_spacing(candidate.get("note", "")),
                "source": candidate["source"],
            }
        )
    return functions[:60]


def extract_acceptance_criteria(doc: ParsedDocument) -> list[dict[str, Any]]:
    criteria: list[dict[str, Any]] = []
    seen: set[str] = set()
    for test_case in extract_test_cases(doc):
        text = test_case.get("acceptance_criteria", "").strip()
        if not text:
            continue
        key = f"{test_case['id']}::{text}"
        if key in seen:
            continue
        seen.add(key)
        criteria.append(
            {
                "id": test_case["id"],
                "name": test_case["name"],
                "criteria": text,
                "source": test_case["source"],
            }
        )
    return criteria[:80]


def schema_bucket_for_metric(metric: dict[str, Any]) -> str:
    section = metric.get("section", "")
    name = metric.get("name", "")
    if section == "orbit_environment":
        return "orbit_environment"
    if section == "electrical_indicators":
        return "electrical_indicators"
    if section == "solar_array":
        return "solar_array_requirements"
    if section == "battery_energy_storage":
        return "battery_requirements"
    if section == "control_equipment":
        return "control_equipment_requirements"
    if section == "reliability_maintenance":
        if "可靠度" in name:
            return "reliability_requirements"
        return "lifetime_requirements"
    if section == "mechanical_constraints":
        if "功耗" in name:
            return "power_budget"
        return "mass_constraints"
    return "misc_requirements"


def group_metrics_for_schema(metrics: list[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
    grouped: dict[str, list[dict[str, Any]]] = {
        "orbit_environment": [],
        "electrical_indicators": [],
        "solar_array_requirements": [],
        "battery_requirements": [],
        "control_equipment_requirements": [],
        "lifetime_requirements": [],
        "maintenance_requirements": [],
        "mass_constraints": [],
        "power_budget": [],
        "reliability_requirements": [],
        "misc_requirements": [],
    }
    for metric in metrics:
        bucket = schema_bucket_for_metric(metric)
        grouped.setdefault(bucket, []).append(metric)
        if metric.get("name", "").startswith("在轨更换/"):
            grouped["maintenance_requirements"].append(metric)
    return grouped


def extract_topology_scheme(doc: ParsedDocument) -> dict[str, Any]:
    features = []
    topology_contexts = topology_context_pages(doc)
    context_text = "\n".join(item["text"] for item in topology_contexts)
    search_space = context_text or doc.text
    for feature_id, (pattern, label) in TOPOLOGY_FEATURES.items():
        match = re.search(pattern, search_space, flags=re.IGNORECASE)
        if not match:
            continue
        source_page = None
        source_text = search_space[max(0, match.start() - 90) : min(len(search_space), match.end() + 120)]
        for item in topology_contexts:
            page_match = re.search(pattern, item["text"], flags=re.IGNORECASE)
            if not page_match:
                continue
            source_page = item["page"]
            source_text = item["text"][max(0, page_match.start() - 90) : min(len(item["text"]), page_match.end() + 120)]
            break
        features.append(
            {
                "id": feature_id,
                "label": label,
                "summary": summarize_topology_feature(feature_id, source_text),
                "source": build_source(source_page, source_text),
            }
        )
    sections = extract_system_sections(doc)
    topology_blocks = sections.get("topology_scheme", [])
    architecture = "未识别"
    channel_count = None
    bus_voltage = None
    regulation_modes: list[str] = []
    if topology_contexts or topology_blocks:
        joined = context_text or "\n".join(
            item["content"]
            for item in topology_blocks
            if not any(marker in item["content"] for marker in TOPOLOGY_NOISE_MARKERS)
        )
        if re.search(r"(两个|2个).{0,8}(功率通道)", joined):
            channel_count = 2
        bus_match = re.search(
            r"(母线电压|一次母线)[^\n]{0,40}?(\d+\s*V\s*(?:-|~|～|至)\s*\d+\s*V)",
            joined,
            flags=re.IGNORECASE,
        )
        if bus_match:
            bus_voltage = normalize_space(bus_match.group(2))
        if re.search(r"S3R|分流调节", joined, flags=re.IGNORECASE):
            regulation_modes.append("光照区S3R分流调节")
        if re.search(r"升压.*放电|放电调节.*升压|Boost|BDR", joined, flags=re.IGNORECASE):
            regulation_modes.append("阴影区升压放电调节")
        summary_parts = []
        if re.search(r"光伏电源系统|太阳电池翼", joined):
            summary_parts.append("光伏电源系统")
        if channel_count:
            summary_parts.append(f"{channel_count}个独立功率通道")
        if re.search(r"四个相对独立的[“\"]?机组|四个机组", joined):
            summary_parts.append("每通道4机组并联冗余")
        if re.search(r"一次母线全调节|母线全调节", joined):
            summary_parts.append("一次母线全调节")
        if regulation_modes:
            summary_parts.extend(regulation_modes[:2])
        if bus_voltage:
            summary_parts.append(f"母线范围{bus_voltage}")
        if summary_parts:
            architecture = "，".join(summary_parts)
    return {
        "architecture": architecture if architecture != "未识别" else " / ".join(feature["label"] for feature in features[:4]) or "未识别",
        "features": features,
        "channel_count": channel_count,
        "bus_voltage_range": bus_voltage,
        "regulation_modes": regulation_modes,
        "context_pages": [item["page"] for item in topology_contexts if item.get("page")],
        "confidence": confidence_from_hits(bool(features), len(features) >= 3, len(features) >= 5),
    }


def extract_design_constraints(doc: ParsedDocument, metrics: list[dict[str, Any]]) -> list[dict[str, Any]]:
    constraints = []
    seen: set[str] = set()
    for metric in metrics:
        if metric["section"] == "mission_function":
            continue
        if metric.get("value", "").strip() in {"—", ""}:
            continue
        text = f"{metric['name']} {metric['value']}"
        if not re.search(r"(≥|≤|不小于|不大于|范围|约束|要求|稳定|恢复|寿命|满足|支持|保障)", text):
            continue
        key = f"{metric['section']}::{metric['name']}::{metric['value']}"
        if key in seen:
            continue
        seen.add(key)
        constraints.append(
            {
                "category": metric["section"],
                "name": metric["name"],
                "requirement": metric["value"],
                "rationale": infer_constraint_rationale(metric["section"], text),
                "source": metric["source"],
            }
        )
    for candidate in extract_table_candidates(doc):
        label = candidate["label"]
        if candidate["section"] == "mission_function":
            continue
        text = f"{label} {candidate['value']}"
        if not is_rule_sentence(text):
            continue
        if not re.search(r"(≥|≤|不小于|不大于|不得|必须|应|要求|支持|满足|寿命|恢复时间|纹波|可靠度|次数)", text):
            continue
        key = f"{candidate['section']}::{label}::{candidate['value']}"
        if key in seen:
            continue
        seen.add(key)
        constraints.append(
            {
                "category": candidate["section"],
                "name": label,
                "requirement": candidate["value"],
                "rationale": infer_constraint_rationale(candidate["section"], text),
                "source": candidate["source"],
            }
        )
    return constraints[:120]


def infer_constraint_rationale(section: str, text: str) -> str:
    if section == "orbit_environment":
        return "作为发电、储能容量和热/光照边界的输入条件。"
    if section == "electrical_indicators":
        return "约束母线供电能力、电能质量和瞬态响应。"
    if section == "solar_array":
        return "约束太阳翼发电能力、展开机构和姿态跟踪能力。"
    if section == "battery_energy_storage":
        return "约束储能容量、放电深度和寿命裕度。"
    if section == "reliability_maintenance":
        return "约束长寿命任务、故障处理和在轨维护能力。"
    if section == "mechanical_constraints":
        return "约束结构动力学、磁特性和系统重量。"
    if "母线" in text:
        return "约束一次母线和特殊负载母线的稳定供电。"
    return "从原文指标表抽取的设计边界。"


def extract_deliverables(doc: ParsedDocument) -> list[dict[str, Any]]:
    deliverables = []
    seen: set[str] = set()
    for candidate in extract_table_candidates(doc):
        if not is_deliverable_candidate(candidate):
            continue
        cells = candidate["cells"]
        row_text = candidate["row_text"]
        if len(cells) >= 4:
            category = cells[1]
            file_name = cells[2]
            file_type = cells[3]
        elif len(cells) >= 3:
            category = cells[0]
            file_name = cells[1]
            file_type = cells[2]
        else:
            category = cells[0]
            file_name = cells[1] if len(cells) > 1 else cells[0]
            file_type = "文件"
        key = f"{category}::{file_name}::{file_type}"
        if key in seen:
            continue
        seen.add(key)
        deliverables.append(
            {
                "category": category,
                "name": file_name,
                "type": file_type,
                "source": make_source(doc, row_text),
            }
        )
    return deliverables[:80]


def build_power_system_model(doc: ParsedDocument) -> dict[str, Any]:
    sections = extract_system_sections(doc)
    functional_requirements = extract_functional_requirements(doc)
    metrics = extract_structured_metrics(doc)
    grouped_metrics = group_metrics_for_schema(metrics)
    topology = extract_topology_scheme(doc)
    constraints = extract_design_constraints(doc, metrics)
    deliverables = extract_deliverables(doc)
    test_cases = extract_test_cases(doc)
    acceptance_criteria = extract_acceptance_criteria(doc)
    section_summary = []
    for section, items in sections.items():
        preview_items = [item["content"][:120] for item in items[:3]]
        section_summary.append(
            {
                "id": section,
                "label": SYSTEM_SECTION_LABELS.get(section, section),
                "items": len(items),
                "preview": preview_items,
            }
        )
    return {
        "schema": "aerospace_power_system.v1",
        "extraction_mode": "docling_schema_pipeline_v1",
        "sections": sections,
        "section_summary": section_summary,
        "functional_requirements": functional_requirements,
        "topology_scheme": topology,
        "metrics": metrics,
        "schema_groups": grouped_metrics,
        "orbit_environment": grouped_metrics["orbit_environment"],
        "electrical_indicators": grouped_metrics["electrical_indicators"],
        "solar_array_requirements": grouped_metrics["solar_array_requirements"],
        "battery_requirements": grouped_metrics["battery_requirements"],
        "control_equipment_requirements": grouped_metrics["control_equipment_requirements"],
        "lifetime_requirements": grouped_metrics["lifetime_requirements"],
        "maintenance_requirements": grouped_metrics["maintenance_requirements"],
        "mass_constraints": grouped_metrics["mass_constraints"],
        "power_budget": grouped_metrics["power_budget"],
        "reliability_requirements": grouped_metrics["reliability_requirements"],
        "constraints": constraints,
        "deliverables": deliverables,
        "test_cases": test_cases,
        "acceptance_criteria": acceptance_criteria,
    }


def extract_topologies(doc: ParsedDocument) -> list[dict[str, Any]]:
    results = []
    seen_ids: set[str] = set()
    for name, pattern in TOPOLOGY_PATTERNS.items():
        matches = list(re.finditer(pattern, doc.text, flags=re.IGNORECASE))
        if matches:
            if name in seen_ids:
                continue
            seen_ids.add(name)
            first = matches[0]
            start = max(0, first.start() - 180)
            end = min(len(doc.text), first.end() + 240)
            snippet = normalize_space(doc.text[start:end])
            source = source_for_text(doc.pages, snippet)
            results.append(
                {
                    "id": name,
                    "label": name.replace("_", " ").title(),
                    "mentions": len(matches),
                    "confidence": confidence_from_hits(len(matches) > 2, True),
                    "source": source,
                }
            )
    return sorted(results, key=lambda item: item["mentions"], reverse=True)[:16]


def extract_parameters(doc: ParsedDocument) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    seen: set[tuple[str, str, str]] = set()
    text = doc.text
    for cells in split_table_rows(text):
        label = cells[0]
        value_text = " ".join(cells[1:])
        if not re.search(r"\d", value_text):
            continue
        for match in re.finditer(rf"(?P<value>{VALUE_PATTERN})\s*(?P<unit>{UNIT_PATTERN})", value_text, flags=re.IGNORECASE):
            canonical = parameter_name_from_label(label)
            value = re.sub(r"\s+", "", match.group("value"))
            unit = match.group("unit")
            key = (canonical, value, unit)
            if key in seen:
                continue
            seen.add(key)
            row_text = " | ".join(cells)
            results.append(
                {
                    "name": canonical,
                    "raw_label": label,
                    "value": value,
                    "unit": unit,
                    "confidence": confidence_from_hits(True, "|" in row_text),
                    "source": source_for_text(doc.pages, row_text),
                }
            )
    for canonical, aliases in PARAM_ALIASES.items():
        alias_pattern = "|".join(re.escape(alias) for alias in aliases)
        patterns = [
            rf"(?P<label>{alias_pattern})\s*[:=]?\s*(?P<value>{VALUE_PATTERN})\s*(?P<unit>{UNIT_PATTERN})",
            rf"(?P<value>{VALUE_PATTERN})\s*(?P<unit>{UNIT_PATTERN})\s+(?P<label>{alias_pattern})",
        ]
        for pattern in patterns:
            for match in re.finditer(pattern, text, flags=re.IGNORECASE):
                label = match.group("label")
                value = re.sub(r"\s+", "", match.group("value"))
                unit = match.group("unit")
                key = (canonical, value, unit)
                if key in seen:
                    continue
                seen.add(key)
                start = max(0, match.start() - 150)
                end = min(len(text), match.end() + 180)
                snippet = normalize_space(text[start:end])
                results.append(
                    {
                        "name": canonical,
                        "raw_label": label,
                        "value": value,
                        "unit": unit,
                        "confidence": confidence_from_hits(True, len(value) > 0),
                        "source": source_for_text(doc.pages, snippet),
                    }
                )
    return results[:80]


def extract_formulas(doc: ParsedDocument) -> list[dict[str, Any]]:
    formulas = []
    formula_pattern = re.compile(
        r"(?P<expr>[A-Za-zΔ∆ηµμ][A-Za-z0-9_Δ∆ηµμ/().\s-]{0,30}\s*=\s*[^。\n;]{4,120})"
    )
    for match in formula_pattern.finditer(doc.text):
        expr = normalize_space(match.group("expr"))
        if len(expr) < 8 or len(re.findall(r"[A-Za-z]", expr)) < 2:
            continue
        variables = sorted(set(re.findall(r"\b[A-Za-z][A-Za-z0-9_]{0,8}\b|[Δ∆][A-Za-z0-9_]+", expr)))
        snippet = normalize_space(doc.text[max(0, match.start() - 160) : min(len(doc.text), match.end() + 180)])
        formulas.append(
            {
                "name": variables[0].lower() if variables else "formula",
                "expression": expr,
                "variables": variables[:12],
                "confidence": confidence_from_hits("=" in expr, any(op in expr for op in ["/", "*", "+", "-"])),
                "source": source_for_text(doc.pages, snippet),
            }
        )
    unique = []
    seen = set()
    for formula in formulas:
        key = formula["expression"].lower()
        if key not in seen:
            seen.add(key)
            unique.append(formula)
    return unique[:60]


def extract_components(doc: ParsedDocument) -> list[dict[str, Any]]:
    components = []
    for component, pattern in COMPONENT_PATTERNS.items():
        matches = list(re.finditer(pattern, doc.text, flags=re.IGNORECASE))
        if not matches:
            continue
        nearby_rules = []
        for sentence in split_sentences(doc.text):
            if re.search(pattern, sentence, flags=re.IGNORECASE) and re.search(
                RULE_CATEGORIES["component_selection"], sentence, flags=re.IGNORECASE
            ):
                nearby_rules.append(sentence[:260])
        first = matches[0]
        snippet = normalize_space(doc.text[max(0, first.start() - 180) : min(len(doc.text), first.end() + 220)])
        components.append(
            {
                "type": component,
                "mentions": len(matches),
                "selection_notes": nearby_rules[:4],
                "confidence": confidence_from_hits(len(matches) > 1, bool(nearby_rules)),
                "source": source_for_text(doc.pages, snippet),
            }
        )
    return sorted(components, key=lambda item: item["mentions"], reverse=True)


def extract_rules(doc: ParsedDocument) -> list[dict[str, Any]]:
    rules = []
    seen: set[str] = set()
    imperative = re.compile(
        r"\b(should|must|keep|place|connect|route|avoid|ensure|use|select|choose|recommend|注意|必须|应|应该|避免|保持|放置|选择|采用|配置|实现|维持|保障|支持|满足|需|要求)\b",
        re.IGNORECASE,
    )
    for candidate in extract_table_candidates(doc):
        row_text = candidate["row_text"]
        if is_deliverable_candidate(candidate):
            continue
        if not is_rule_sentence(row_text):
            continue
        if is_test_case_text(row_text):
            row_text = candidate["value"]
            if not is_rule_sentence(row_text):
                continue
        matched_categories = [
            category
            for category, pattern in RULE_CATEGORIES.items()
            if re.search(pattern, row_text, flags=re.IGNORECASE)
        ]
        if not matched_categories:
            continue
        if not re.search(r"(≥|≤|必须|应|要求|禁止|不得|支持|满足|恢复时间|纹波|寿命|可靠度|次数|能力)", row_text):
            continue
        key = re.sub(r"\W+", "", row_text.lower())[:160]
        if key in seen:
            continue
        seen.add(key)
        severity = (
            "high"
            if any(category in matched_categories for category in ["bus_regulation", "power_quality", "protection", "reliability_maintenance"])
            else "medium"
        )
        rules.append(
            {
                "category": matched_categories[0],
                "rule": row_text[:420],
                "severity": severity,
                "confidence": confidence_from_hits(True, len(candidate["cells"]) > 2),
                "source": source_for_text(doc.pages, row_text),
            }
        )
    for sentence in split_sentences(doc.text):
        if not is_rule_sentence(sentence):
            continue
        if is_test_case_text(sentence):
            continue
        matched_categories = [
            category
            for category, pattern in RULE_CATEGORIES.items()
            if re.search(pattern, sentence, flags=re.IGNORECASE)
        ]
        if not matched_categories:
            continue
        has_action = bool(imperative.search(sentence))
        if not has_action and not re.search(r"(≥|≤|不得|禁止|寿命|可靠度|范围内|恢复时间|纹波|约束)", sentence):
            continue
        key = re.sub(r"\W+", "", sentence.lower())[:160]
        if key in seen:
            continue
        seen.add(key)
        severity = (
            "high"
            if any(category in matched_categories for category in ["layout", "emi", "protection", "bus_regulation", "power_quality", "reliability_maintenance"])
            else "medium"
        )
        rules.append(
            {
                "category": matched_categories[0],
                "rule": sentence[:420],
                "severity": severity,
                "confidence": confidence_from_hits(has_action, len(matched_categories) > 1),
                "source": source_for_text(doc.pages, sentence),
            }
        )
    return rules[:120]


def build_summary(
    doc: ParsedDocument,
    topologies: list[dict[str, Any]],
    rules: list[dict[str, Any]],
    power_system: dict[str, Any] | None = None,
) -> dict[str, Any]:
    dominant_rules = Counter(rule["category"] for rule in rules).most_common(4)
    topology_scheme = (power_system or {}).get("topology_scheme", {})
    metrics = (power_system or {}).get("metrics", [])
    constraints = (power_system or {}).get("constraints", [])
    functions = (power_system or {}).get("functional_requirements", [])
    test_cases = (power_system or {}).get("test_cases", [])
    feature_tags = [item["label"] for item in topology_scheme.get("features", [])[:6]]
    metric_tags = []
    for item in metrics:
        name = item.get("name", "")
        if not name or len(name) > 10:
            continue
        if name in metric_tags:
            continue
        metric_tags.append(name)
        if len(metric_tags) >= 6:
            break
    top_keywords = feature_tags + metric_tags
    digest = []
    if topology_scheme.get("architecture"):
        digest.append(topology_scheme["architecture"])
    if functions:
        digest.append(f"识别 {len(functions)} 项功能定义")
    if metrics:
        digest.append(f"已归并 {len(metrics)} 条结构化指标")
    if constraints:
        digest.append(f"提炼 {len(constraints)} 条设计约束")
    if test_cases:
        digest.append(f"抽取 {len(test_cases)} 条测试验证项")
    return {
        "primary_topology": topologies[0]["id"] if topologies else None,
        "estimated_tokens": math.ceil(len(doc.text) / 4),
        "character_count": len(doc.text),
        "keyword_cloud": top_keywords,
        "digest": digest,
        "dominant_rule_categories": [{"category": category, "count": count} for category, count in dominant_rules],
    }


def extract_knowledge(doc: ParsedDocument) -> dict[str, Any]:
    topologies = extract_topologies(doc)
    parameters = extract_parameters(doc)
    formulas = extract_formulas(doc)
    components = extract_components(doc)
    rules = extract_rules(doc)
    power_system = build_power_system_model(doc)
    return {
        "schema_version": SCHEMA_VERSION,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "document": {
            "title": Path(doc.name).stem,
            "filename": doc.name,
            "type": doc.file_type,
            "pages": len(doc.pages),
        },
        "summary": build_summary(doc, topologies, rules, power_system),
        "topologies": topologies,
        "parameters": parameters,
        "formulas": formulas,
        "components": components,
        "design_rules": rules,
        "test_cases": power_system.get("test_cases", []),
        "power_system": power_system,
    }


def page_signal_counts(payload: dict[str, Any], page_count: int) -> list[dict[str, int]]:
    rows = [{"page": index, "parameters": 0, "formulas": 0, "rules": 0, "components": 0} for index in range(1, page_count + 1)]
    buckets = [
        ("parameters", payload["parameters"]),
        ("formulas", payload["formulas"]),
        ("rules", payload["design_rules"]),
        ("components", payload["components"]),
    ]
    for key, items in buckets:
        for item in items:
            page = item.get("source", {}).get("page")
            if isinstance(page, int) and 1 <= page <= page_count:
                rows[page - 1][key] += 1
    return rows


def asset_data_url(path: Path, mime_type: str) -> str:
    if not path.exists():
        return ""
    encoded = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:{mime_type};base64,{encoded}"


def render_metric(label: str, value: Any) -> None:
    display_value = value
    if label in {"主拓扑", "系统架构摘要"} and isinstance(value, str):
        display_value = TOPOLOGY_LABELS.get(value, value)
    st.markdown(
        f"""
        <div class="metric-tile">
          <span>{escape(str(label))}</span>
          <strong>{escape(str(display_value))}</strong>
        </div>
        """,
        unsafe_allow_html=True,
    )


def slugify_label(text: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")
    return slug or "item"


def section_label(title: str, caption: str | None = None) -> None:
    caption_html = f"<p>{escape(caption)}</p>" if caption else ""
    st.markdown(
        f"""
        <div class="section-label">
          <h2>{escape(title)}</h2>
          {caption_html}
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_empty_state() -> None:
    st.markdown(
        """
        <div class="empty-workbench">
          <div>
            <span class="system-kicker">准备开始解析</span>
            <h2>上传航天电源设计文档</h2>
            <p>PowerDoc-KB 会把系统级设计文本整理成拓扑方案、关键指标、设计约束、交付清单和可追溯 JSON。</p>
          </div>
          <div class="empty-grid">
            <div><strong>1</strong><span>上传 PDF 或 Word</span></div>
            <div><strong>2</strong><span>生成电源系统结构</span></div>
            <div><strong>3</strong><span>查看分区与约束</span></div>
            <div><strong>4</strong><span>导出知识 JSON</span></div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_source(source: dict[str, Any]) -> str:
    page = source.get("page")
    page_text = f"P{page}" if page else "原文"
    snippet = escape(source.get("snippet", ""))
    if snippet:
        return f'<div class="source-inline"><span>{page_text}</span><span>{snippet}</span></div>'
    return f'<div class="source-inline"><span>{page_text}</span></div>'


def card(title: str, body: str, meta: str = "") -> None:
    st.markdown(
        f"""
        <div class="kb-card">
          <div class="kb-card-title">{escape(title)}</div>
          <div class="kb-card-body">{body}</div>
          <div class="kb-card-meta">{meta}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def code_panel(text: str, max_height: int = 560) -> None:
    st.markdown(
        f"""
        <div class="code-panel" style="max-height:{max_height}px;"><pre><code>{escape(text)}</code></pre></div>
        """,
        unsafe_allow_html=True,
    )


def render_dark_table(headers: list[str], rows: list[list[Any]]) -> None:
    head_html = "".join(f"<th>{escape(header)}</th>" for header in headers)
    body_html = "".join(
        "<tr>" + "".join(f"<td>{escape('' if value is None else str(value))}</td>" for value in row) + "</tr>"
        for row in rows
    )
    st.markdown(
        f"""
        <div class="table-shell">
          <table class="dark-table">
            <thead><tr>{head_html}</tr></thead>
            <tbody>{body_html}</tbody>
          </table>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_document_switcher(documents: list[dict[str, Any]], active_index: int) -> int:
    if len(documents) <= 1:
        return 0
    st.markdown(
        """
        <div class="selector-shell">
          <strong>当前文档</strong>
          <span>已载入的文档会保留在本地会话中，可直接切换查看结果。</span>
        </div>
        """,
        unsafe_allow_html=True,
    )
    cols = st.columns(len(documents))
    selected = active_index
    for idx, (column, item) in enumerate(zip(cols, documents)):
        label = f"{item['name']} · P{item['pages']} · {item['chars']}字"
        button_type = "primary" if idx == active_index else "secondary"
        with column:
            if st.button(
                label,
                key=f"powerdoc_active_doc_btn_{idx}",
                use_container_width=True,
                type=button_type,
            ):
                selected = idx
    return selected


def render_current_document_caption(document: dict[str, Any]) -> None:
    st.markdown(
        f"""
        <div class="current-doc-caption">
          当前文档：{escape(document["name"])} / {escape(str(document["pages"]))} 页或段 / {escape(str(document["chars"]))} 字符
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_parse_done_banner(count: int) -> None:
    st.markdown(
        f"""
        <div class="parse-done-banner">
          <strong>解析完成</strong>
          <span>当前已载入 {count} 份文档，可直接切换视图或导出 JSON。</span>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_workspace_masthead(video_src: str, has_documents: bool, document_count: int = 0) -> None:
    if has_documents:
        st.markdown(
            f"""
            <div class="hero-band hero-band-compact">
              <div class="hero-content hero-content-compact">
                <div class="hero-copy">
                  <span class="system-kicker">卫星空间电源知识建模</span>
                  <h1>PowerDoc-KB</h1>
                  <p>围绕当前文档队列，继续浏览系统模型、设计约束与可追溯 JSON。</p>
                  <div class="hero-status hero-status-compact">
                    <div><span>当前队列</span><strong>{document_count} 份文档</strong></div>
                    <div><span>工作模式</span><strong>本地解析工作台</strong></div>
                    <div><span>结果形态</span><strong>结构化知识 JSON</strong></div>
                  </div>
                </div>
                <div class="video-stage video-stage-compact">
                  <video class="hero-video" autoplay muted loop playsinline preload="auto" src="{video_src}"></video>
                  <div class="video-glass-line"></div>
                </div>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        return

    st.markdown(
        f"""
        <div class="hero-band">
          <div class="hero-content">
            <div class="hero-copy">
              <span class="system-kicker">卫星空间电源知识建模</span>
              <h1>PowerDoc-KB</h1>
              <p>将航天电源设计文档整理成系统拓扑、关键指标、设计约束和可追溯 JSON。</p>
              <div class="mission-chips">
                <span>多文件导入</span>
                <span>系统结构化</span>
                <span>可追溯抽取</span>
              </div>
              <div class="hero-status">
                <div><span>输入</span><strong>PDF / DOCX / TXT / MD</strong></div>
                <div><span>输出</span><strong>电源系统知识 JSON</strong></div>
              </div>
            </div>
            <div class="video-stage">
              <video class="hero-video" autoplay muted loop playsinline preload="auto" src="{video_src}"></video>
              <div class="video-glass-line"></div>
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_result_header(document: dict[str, Any], count: int, payload: dict[str, Any]) -> None:
    power_system = payload.get("power_system", {})
    metrics_count = len(power_system.get("metrics", []))
    constraints_count = len(power_system.get("constraints", []))
    deliverables_count = len(power_system.get("deliverables", []))
    topology_label = TOPOLOGY_LABELS.get(
        payload.get("summary", {}).get("primary_topology"),
        payload.get("summary", {}).get("primary_topology") or "未识别",
    )
    st.markdown(
        f"""
        <div class="result-header-shell">
          <div class="result-header-copy">
            <strong>解析完成</strong>
            <span>{count} 份文档已载入，当前聚焦 <em>{escape(document["name"])}</em></span>
          </div>
          <div class="result-header-grid">
            <div><span>当前文档</span><strong>{escape(str(document["pages"]))} 页 / {escape(str(document["chars"]))} 字</strong></div>
            <div><span>主拓扑</span><strong>{escape(str(topology_label))}</strong></div>
            <div><span>结构化指标</span><strong>{metrics_count} 项</strong></div>
            <div><span>设计约束</span><strong>{constraints_count} 项</strong></div>
            <div><span>交付文件</span><strong>{deliverables_count} 项</strong></div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_processing_banner(lines: list[str]) -> None:
    items = "".join(f"<li>{escape(line)}</li>" for line in lines)
    st.markdown(
        f"""
        <div class="parse-progress-shell">
          <strong>正在解析文档</strong>
          <ul>{items}</ul>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_sidebar_empty_state() -> None:
    st.markdown(
        """
        <div class="sidebar-guide-card">
          <strong>等待文档</strong>
          <span>请在右侧工作区添加 PDF、Word 或文本文件。</span>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_sidebar_active_document(document: dict[str, Any]) -> None:
    st.markdown(
        f"""
        <div class="sidebar-active-card">
          <div class="sidebar-active-icon">文</div>
          <div class="sidebar-active-meta">
            <strong>{escape(document["name"])}</strong>
            <span>P{escape(str(document["pages"]))} / {escape(str(document["chars"]))} 字</span>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_sidebar_file_list(documents: list[dict[str, Any]], active_index: int = 0) -> None:
    if not documents:
        return
    if len(documents) <= 1:
        return
    cards = ""
    for index, item in enumerate(documents):
        active_class = " is-active" if index == active_index else ""
        cards += f"""
        <div class="upload-file-card{active_class}">
          <div class="upload-file-icon">{index + 1}</div>
          <div class="upload-file-meta">
            <strong>{escape(item['name'])}</strong>
            <span>P{item['pages']} / {item['chars']} 字</span>
          </div>
        </div>
        """
    st.markdown(f'<div class="upload-file-list">{cards}</div>', unsafe_allow_html=True)


def reset_uploaded_files() -> None:
    st.session_state["powerdoc_parsed_items"] = []
    st.session_state["powerdoc_show_uploader"] = True
    st.session_state["powerdoc_active_doc_index"] = 0
    st.session_state["powerdoc_active_view"] = "System Model"
    st.session_state["powerdoc_active_section"] = ""
    st.session_state["powerdoc_uploader_version"] = st.session_state.get("powerdoc_uploader_version", 0) + 1


def render_upload_panel(show_uploader: bool, file_count: int) -> list[Any] | None:
    has_files = file_count > 0
    uploader_version = st.session_state.get("powerdoc_uploader_version", 0)
    if not show_uploader and has_files:
        st.markdown(
            f"""
            <div class="upload-surface upload-surface-loaded">
              <div class="upload-surface-head">
                <div class="upload-surface-copy">
                  <strong>文档队列已就绪</strong>
                  <span>当前会话已载入 {file_count} 份文档，可直接切换视图查看结构化结果。</span>
                </div>
                <div class="upload-surface-badge">本地会话</div>
              </div>
              <div class="upload-loaded-shell">
                <div class="upload-loaded-meta">
                  <strong>已载入 {file_count} 份文档</strong>
                  <span>上传区已自动收起，仅保留结果浏览与 JSON 导出能力。</span>
                </div>
              </div>
              <div class="upload-surface-foot">
                <span>支持多文件批量解析</span>
                <span>保留原文溯源信息</span>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.button("重新选择文件", key="powerdoc_repick", on_click=reset_uploaded_files, use_container_width=False)
        return None

    title = "更换文档" if has_files else "添加文档"
    hint = "已载入文档，可继续选择新文件替换当前队列。" if has_files else "选择 PDF、Word 或文本文件，支持一次上传多个。"
    st.markdown(
        f"""
        <div class="upload-surface">
          <div class="upload-surface-head">
            <div class="upload-surface-copy">
              <strong>{title}</strong>
              <span>{hint}</span>
            </div>
            <div class="upload-surface-badge">本地解析</div>
          </div>
          <div class="upload-surface-metrics">
            <div><span>支持格式</span><strong>PDF / DOCX / TXT / MD</strong></div>
            <div><span>上传方式</span><strong>单次可选多个文件</strong></div>
            <div><span>结果输出</span><strong>结构化知识 JSON</strong></div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    uploader_host = st.empty()
    with uploader_host.container():
        uploaded_files = st.file_uploader(
            "选择电源设计文档",
            type=["pdf", "docx", "txt", "md"],
            accept_multiple_files=True,
            key=f"powerdoc_uploads_{uploader_version}",
            label_visibility="collapsed",
            help="本地解析，不会上传到外部服务。",
        )
    if uploaded_files:
        uploader_host.empty()
    st.markdown(
        """
        <div class="upload-surface-foot">
          <span>单个文件上限 200MB</span>
          <span>上传后自动开始解析</span>
          <span>处理过程仅保留在当前本地会话</span>
        </div>
        """,
        unsafe_allow_html=True,
    )
    return uploaded_files


def render_view_selector(enabled_views: list[str]) -> str:
    labels = {
        "System Model": "系统模型",
        "Section Browser": "分区浏览",
        "Visual Parse": "解析摘要",
        "Overview": "总览",
        "Parameters": "参数",
        "Formulas": "公式",
        "Components": "器件",
        "Design Rules": "规则",
        "JSON": "JSON",
    }
    active = st.session_state.get("powerdoc_active_view", enabled_views[0])
    if active not in enabled_views:
        active = enabled_views[0]
    st.markdown(
        """
        <div class="selector-shell">
          <strong>视图切换</strong>
          <span>围绕同一份结构化结果切换系统模型、摘要和 JSON 输出。</span>
        </div>
        """,
        unsafe_allow_html=True,
    )
    cols = st.columns(len(enabled_views))
    selected = active
    for view, column in zip(enabled_views, cols):
        button_type = "primary" if view == active else "secondary"
        with column:
            if st.button(
                labels.get(view, view),
                key=f"powerdoc_view_btn_{view}",
                use_container_width=True,
                type=button_type,
            ):
                selected = view
    st.session_state["powerdoc_active_view"] = selected
    return selected


def render_section_selector(section_labels: list[str], active_label: str) -> str:
    if not section_labels:
        return active_label
    selected_key = st.session_state.get("powerdoc_active_section", active_label)
    if selected_key not in section_labels:
        selected_key = active_label
    st.markdown(
        """
        <div class="selector-shell selector-shell-tight">
          <strong>分区浏览</strong>
          <span>按系统分区查看原文片段与溯源页码。</span>
        </div>
        """,
        unsafe_allow_html=True,
    )
    selected = selected_key
    rows = [section_labels[i : i + 4] for i in range(0, len(section_labels), 4)]
    for row_idx, row_labels in enumerate(rows):
        cols = st.columns(len(row_labels))
        for col_idx, (column, label) in enumerate(zip(cols, row_labels)):
            button_type = "primary" if label == selected_key else "secondary"
            with column:
                if st.button(
                    label,
                    key=f"powerdoc_section_btn_{row_idx}_{col_idx}",
                    use_container_width=True,
                    type=button_type,
                ):
                    selected = label
    st.session_state["powerdoc_active_section"] = selected
    return selected


def render_parse_visualization(payload: dict[str, Any], doc: ParsedDocument) -> None:
    section_label("解析流程", "展示当前文档从读取到结构化入库的主要处理阶段。")
    total_items = (
        len(payload["parameters"])
        + len(payload["formulas"])
        + len(payload["components"])
        + len(payload["design_rules"])
        + len(payload["topologies"])
    )
    pipeline = [
        ("01", "文档读取", f"{payload['document']['pages']} 页或段"),
        ("02", "文本切分", f"{len(split_sentences(doc.text))} 个文本片段"),
        ("03", "信息抽取", f"{total_items} 个抽取结果"),
        ("04", "原文绑定", "保留页码与片段"),
        ("05", "结构化输出", payload["schema_version"]),
    ]
    st.markdown('<div class="flow-lane-wrap">', unsafe_allow_html=True)
    for column, (index, title, detail) in zip(st.columns(len(pipeline)), pipeline):
        with column:
            st.markdown(
                f"""
                <div class="flow-step">
                  <span>{escape(index)}</span>
                  <strong>{escape(title)}</strong>
                  <small>{escape(detail)}</small>
                </div>
                """,
                unsafe_allow_html=True,
            )
    st.markdown("</div>", unsafe_allow_html=True)

    section_label("解析结果构成", "说明当前文档主要抽出了哪些内容，避免无意义的视觉噪音。")
    rows = [
        ["系统拓扑", len(payload["topologies"]), "识别供电架构、调节方式和功能链路"],
        ["结构化参数", len(payload["parameters"]), "提取数值、单位和原始标签"],
        ["器件线索", len(payload["components"]), "识别太阳翼、蓄电池、调节器、滤波器等对象"],
        ["设计规则", len(payload["design_rules"]), "提取约束、要求、寿命和可靠性规则"],
        ["系统模型", len(payload["power_system"].get("metrics", [])), "整理为可入库的系统级字段"],
    ]
    render_dark_table(["内容类型", "数量", "说明"], rows)

    section_label("当前文档重点", "帮助你快速判断这份文档更偏系统方案、指标约束还是交付清单。")
    power_system = payload.get("power_system", {})
    summary_rows = [
        ["主拓扑", TOPOLOGY_LABELS.get(payload["summary"].get("primary_topology"), payload["summary"].get("primary_topology") or "未识别")],
        ["系统分区数", len(power_system.get("sections", {}))],
        ["结构化指标数", len(power_system.get("metrics", []))],
        ["设计约束数", len(power_system.get("constraints", []))],
        ["交付文件数", len(power_system.get("deliverables", []))],
    ]
    st.markdown('<div class="summary-strip">', unsafe_allow_html=True)
    for col, item in zip(st.columns(len(summary_rows)), summary_rows):
        with col:
            st.markdown(
                f"""
                <div class="summary-chip">
                  <span>{escape(str(item[0]))}</span>
                  <strong>{escape(str(item[1]))}</strong>
                </div>
                """,
                unsafe_allow_html=True,
            )
    st.markdown("</div>", unsafe_allow_html=True)

    category_counts = Counter(rule["category"] for rule in payload["design_rules"])
    if category_counts:
        section_label("规则分布", "看出这份文档主要强调哪一类设计边界。")
        bars = ""
        max_count = max(category_counts.values(), default=1)
        for category, count in category_counts.most_common():
            width = max(16, round(count / max_count * 100))
            bars += (
                f'<div class="signal-bar"><div><span>{escape(RULE_CATEGORY_LABELS.get(category, category))}</span><strong>{count}</strong></div>'
                f'<i style="width:{width}%"></i></div>'
            )
        st.markdown(f'<div class="bar-panel">{bars}</div>', unsafe_allow_html=True)


def inject_css() -> None:
    css = """
        <style>
          :root {
            --pd-bg: #030814;
            --pd-surface: rgba(7, 18, 34, .82);
            --pd-surface-muted: rgba(13, 34, 58, .72);
            --pd-ink: #ecf8ff;
            --pd-muted: #8ea8ba;
            --pd-line: rgba(119, 217, 255, .22);
            --pd-accent: #2ee6ff;
            --pd-accent-dark: #7df3ff;
            --pd-green: #ffb84d;
            --pd-radius: 8px;
          }
          .stApp {
            background:
              radial-gradient(circle at 18% 18%, rgba(46, 230, 255, .18), transparent 30%),
              radial-gradient(circle at 78% 10%, rgba(255, 184, 77, .12), transparent 26%),
              radial-gradient(circle at 64% 74%, rgba(62, 92, 255, .12), transparent 32%),
              linear-gradient(180deg, #06101f 0%, #030814 52%, #01040b 100%);
            color: var(--pd-ink);
            font-family: ui-sans-serif, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
          }
          .stApp::before {
            content: "";
            position: fixed;
            inset: 0;
            z-index: 0;
            pointer-events: none;
            background:
              radial-gradient(circle, rgba(255,255,255,.72) 0 1px, transparent 1.6px) 0 0 / 96px 96px,
              radial-gradient(circle, rgba(125,243,255,.45) 0 1px, transparent 1.8px) 38px 24px / 142px 142px;
            opacity: .24;
            animation: star-drift 42s linear infinite;
          }
          .stApp::after {
            content: "";
            position: fixed;
            inset: auto -18% -28% -18%;
            height: 46vh;
            z-index: 0;
            pointer-events: none;
            background: radial-gradient(ellipse at center, rgba(46, 230, 255, .20), transparent 62%);
            filter: blur(34px);
          }
          .block-container {
            position: relative;
            z-index: 1;
            max-width: 1500px;
            padding-top: 22px;
            padding-bottom: 48px;
          }
          [data-testid="stHeader"] {
            background: rgba(3, 8, 20, .72);
            backdrop-filter: blur(12px);
          }
          h1, h2, h3 {
            letter-spacing: 0;
          }
          [data-testid="stSidebar"] {
            background:
              linear-gradient(180deg, rgba(7, 18, 34, .98), rgba(3, 8, 20, .98)),
              radial-gradient(circle at 30% 0%, rgba(46, 230, 255, .14), transparent 34%);
            border-right: 1px solid rgba(125,243,255,.16);
          }
          [data-testid="stSidebar"] * {
            color: rgba(244, 249, 252, .96);
          }
          [data-testid="stSidebar"] h2,
          [data-testid="stSidebar"] h3 {
            color: #ffffff;
            font-size: 16px;
          }
          [data-testid="stSidebar"] p,
          [data-testid="stSidebar"] label {
            color: rgba(226, 238, 246, .86);
          }
          [data-testid="stSidebar"] [data-testid="stElementContainer"],
          [data-testid="stSidebar"] [data-testid="stMarkdown"],
          [data-testid="stSidebar"] [data-testid="stMarkdownContainer"],
          [data-testid="stSidebar"] [data-testid="stVerticalBlock"],
          [data-testid="stSidebar"] [data-testid="stSidebarUserContent"],
          [data-testid="stSidebar"] [data-testid="stHeading"],
          [data-testid="stSidebar"] [data-testid="stCaptionContainer"],
          [data-testid="stSidebar"] .stHeading,
          [data-testid="stSidebar"] .stCaption {
            background: transparent !important;
            box-shadow: none !important;
            border: 0 !important;
          }
          [data-testid="stSidebar"] [data-testid="stFileUploader"],
          [data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"],
          [data-testid="stSidebar"] [data-testid="stFileUploaderFileList"],
          [data-testid="stSidebar"] [data-testid="stFileUploaderFile"] {
            display: none !important;
          }
          [data-testid="stSidebar"] button {
            border-radius: var(--pd-radius);
          }
          [data-testid="stSidebar"] div[data-baseweb="select"] > div,
          [data-testid="stSidebar"] div[data-baseweb="select"] {
            background: rgba(8, 24, 43, .82) !important;
            border: 1px solid rgba(125, 243, 255, .22) !important;
            box-shadow: none !important;
            border-radius: 14px !important;
          }
          .sidebar-brand-shell {
            display: grid;
            gap: 10px;
            padding: 6px 4px 18px;
            margin-bottom: 18px;
            border-bottom: 1px solid rgba(255,255,255,.07);
          }
          .sidebar-brand-shell strong {
            color: #f7fbff;
            font-size: 28px;
            line-height: .95;
            letter-spacing: -.04em;
            font-weight: 780;
          }
          .sidebar-brand-shell span {
            color: #a9c1d1;
            font-size: 13px;
            line-height: 1.55;
            max-width: 18ch;
          }
          .upload-file-list {
            display: none;
          }
          .upload-file-card {
            display: grid;
            grid-template-columns: 34px minmax(0, 1fr);
            gap: 10px;
            align-items: center;
            padding: 10px 11px;
            border-radius: 12px;
            background: rgba(9, 21, 35, .76);
            border: 1px solid rgba(226,249,255,.08);
          }
          .upload-file-card.is-active {
            background: rgba(226,249,255,.10);
            border-color: rgba(125,243,255,.24);
          }
          .upload-file-icon {
            width: 34px;
            height: 34px;
            border-radius: 9px;
            background: rgba(226,249,255,.10);
            border: 1px solid rgba(226,249,255,.10);
            display: grid;
            place-items: center;
            color: #ecf8ff;
            font-size: 12px;
            font-weight: 800;
          }
          .upload-file-meta {
            min-width: 0;
          }
          .upload-file-meta strong {
            display: block;
            color: #f5fbff;
            font-size: 13px;
            line-height: 1.35;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
          }
          .upload-file-meta span {
            display: block;
            margin-top: 4px;
            color: #b8cedc;
            font-size: 11px;
          }
          .sidebar-active-card {
            display: grid;
            grid-template-columns: 42px minmax(0, 1fr);
            gap: 12px;
            align-items: center;
            margin: 10px 0 12px;
            padding: 13px 12px;
            border-radius: 14px;
            background: rgba(10, 24, 40, .94);
            border: 1px solid rgba(125,243,255,.16);
          }
          .sidebar-active-icon {
            width: 42px;
            height: 42px;
            display: grid;
            place-items: center;
            border-radius: 11px;
            background: rgba(226,249,255,.08);
            color: #ecf8ff;
            font-size: 15px;
            font-weight: 800;
          }
          .sidebar-active-meta {
            min-width: 0;
          }
          .sidebar-active-meta strong {
            display: block;
            color: #f7fbff;
            font-size: 13px;
            line-height: 1.35;
            word-break: break-word;
          }
          .sidebar-active-meta span {
            display: block;
            margin-top: 5px;
            color: #d8edf8;
            font-size: 12px;
          }
          .sidebar-guide-card {
            margin-top: 10px;
            padding: 14px 12px;
            border-radius: 14px;
            background: rgba(10, 24, 40, .78);
            border: 1px solid rgba(226,249,255,.10);
          }
          .sidebar-guide-card strong {
            display: block;
            color: #f5fbff;
            font-size: 14px;
            margin-bottom: 6px;
          }
          .sidebar-guide-card span {
            display: block;
            color: #b8cedc;
            font-size: 12px;
            line-height: 1.5;
          }
          .upload-surface {
            margin: -8px 0 0;
            padding: 20px 20px 18px;
            border: 1px solid rgba(226,249,255,.14);
            border-bottom: 0;
            border-radius: 18px 18px 0 0;
            background:
              linear-gradient(180deg, rgba(11, 28, 48, .98), rgba(7, 20, 36, .94));
            box-shadow:
              inset 0 1px 0 rgba(255,255,255,.06),
              0 26px 56px rgba(0, 0, 0, .18);
          }
          .upload-surface-loaded {
            border-bottom: 1px solid rgba(125,243,255,.18);
            border-radius: 18px;
          }
          .upload-surface-head {
            display: flex;
            justify-content: space-between;
            align-items: flex-start;
            gap: 16px;
          }
          .upload-surface-copy {
            min-width: 0;
          }
          .upload-surface-copy strong {
            display: block;
            color: #f5fbff;
            font-size: 18px;
            line-height: 1.2;
            margin-bottom: 7px;
          }
          .upload-surface-copy span {
            display: block;
            color: #dbeef9;
            font-size: 13px;
            line-height: 1.55;
          }
          .upload-surface-badge {
            display: inline-flex;
            align-items: center;
            min-height: 28px;
            padding: 0 12px;
            border-radius: 999px;
            border: 1px solid rgba(125,243,255,.22);
            background: rgba(125,243,255,.08);
            color: #ecf9ff;
            font-size: 12px;
            font-weight: 760;
            white-space: nowrap;
          }
          .upload-surface-metrics {
            display: grid;
            grid-template-columns: repeat(3, minmax(0, 1fr));
            gap: 12px;
            margin-top: 18px;
          }
          .upload-surface-metrics div {
            display: grid;
            gap: 7px;
            min-height: 92px;
            padding: 14px 15px;
            border-radius: 14px;
            border: 1px solid rgba(226,249,255,.10);
            background: linear-gradient(180deg, rgba(226,249,255,.045), rgba(226,249,255,.025));
          }
          .upload-surface-metrics span {
            color: #b8cedc;
            font-size: 12px;
          }
          .upload-surface-metrics strong {
            color: #f3fbff;
            font-size: 16px;
            line-height: 1.45;
          }
          .upload-loaded-shell {
            margin: 16px 0 0;
            padding: 16px 18px;
            border: 1px solid rgba(125,243,255,.18);
            border-radius: 16px;
            background:
              linear-gradient(180deg, rgba(8, 21, 37, .96), rgba(6, 16, 29, .94));
            box-shadow: inset 0 1px 0 rgba(255,255,255,.04);
          }
          .upload-loaded-meta strong {
            display: block;
            color: #f5fbff;
            font-size: 15px;
            line-height: 1.25;
            margin-bottom: 6px;
          }
          .upload-loaded-meta span {
            display: block;
            color: #d7ecf8;
            font-size: 13px;
            line-height: 1.5;
          }
          .upload-surface-foot {
            display: flex;
            flex-wrap: wrap;
            gap: 8px;
            margin: 0 0 20px;
            padding: 11px 0 0;
          }
          .upload-surface-foot span {
            display: inline-flex;
            align-items: center;
            min-height: 28px;
            padding: 0 12px;
            border-radius: 999px;
            background: rgba(226,249,255,.06);
            border: 1px solid rgba(226,249,255,.10);
            color: #d4e9f4;
            font-size: 12px;
            font-weight: 650;
          }
          .block-container [data-testid="stButton"][data-testid*="powerdoc_repick"] {
            margin-top: 14px;
          }
          .selector-shell {
            display: grid;
            gap: 6px;
            margin: 8px 0 12px;
          }
          .selector-shell strong {
            color: #f4fbff;
            font-size: 15px;
            line-height: 1.2;
          }
          .selector-shell span {
            color: #d9eef9;
            font-size: 13px;
            line-height: 1.5;
          }
          .selector-shell-tight {
            margin-top: 4px;
          }
          .parse-progress-shell {
            display: grid;
            gap: 10px;
            margin: 8px 0 18px;
            padding: 14px 16px;
            border-radius: 14px;
            border: 1px solid rgba(125,243,255,.16);
            background: rgba(6, 17, 31, .76);
          }
          .parse-progress-shell strong {
            color: #f4fbff;
            font-size: 15px;
            line-height: 1.2;
          }
          .parse-progress-shell ul {
            margin: 0;
            padding-left: 18px;
          }
          .parse-progress-shell li {
            color: #d9eef9;
            font-size: 13px;
            line-height: 1.6;
            margin: 0;
          }
          .result-header-shell {
            display: grid;
            gap: 16px;
            margin: 8px 0 20px;
            padding: 18px 18px 16px;
            border-radius: 18px;
            border: 1px solid rgba(125,243,255,.16);
            background:
              linear-gradient(180deg, rgba(8, 22, 38, .84), rgba(5, 14, 28, .80));
            box-shadow: 0 24px 54px rgba(0, 0, 0, .18), inset 0 1px 0 rgba(255,255,255,.05);
          }
          .result-header-copy {
            display: grid;
            gap: 6px;
          }
          .result-header-copy strong {
            color: #f4fbff;
            font-size: 18px;
            line-height: 1.15;
          }
          .result-header-copy span {
            color: #d7ebf6;
            font-size: 14px;
            line-height: 1.55;
          }
          .result-header-copy em {
            color: #f6fbff;
            font-style: normal;
            font-weight: 700;
          }
          .result-header-grid {
            display: grid;
            grid-template-columns: repeat(5, minmax(0, 1fr));
            gap: 12px;
          }
          .result-header-grid div {
            display: grid;
            gap: 8px;
            min-height: 82px;
            padding: 13px 14px;
            border-radius: 14px;
            border: 1px solid rgba(226,249,255,.10);
            background: rgba(226,249,255,.035);
          }
          .result-header-grid span {
            color: #a9c2d1;
            font-size: 12px;
          }
          .result-header-grid strong {
            color: #f4fbff;
            font-size: 15px;
            line-height: 1.4;
          }
          .parse-done-banner {
            display: grid;
            gap: 6px;
            margin: 8px 0 18px;
            padding: 14px 16px;
            border-radius: 14px;
            border: 1px solid rgba(125,243,255,.16);
            background: rgba(6, 17, 31, .76);
          }
          .parse-done-banner strong {
            color: #f4fbff;
            font-size: 15px;
            line-height: 1.2;
          }
          .parse-done-banner span {
            color: #d9eef9;
            font-size: 13px;
            line-height: 1.5;
          }
          .current-doc-caption {
            margin: 8px 0 18px;
            color: #eef9ff;
            font-size: 15px;
            line-height: 1.55;
            font-weight: 560;
          }
          .block-container .stButton {
            margin-top: 12px;
            margin-bottom: 18px;
          }
          .block-container button[kind="primary"] {
            background:
              linear-gradient(180deg, rgba(29, 68, 108, .98), rgba(15, 39, 68, .96)) !important;
            color: #f7fdff !important;
            border: 1px solid rgba(125,243,255,.38) !important;
            box-shadow: 0 18px 38px rgba(0, 0, 0, .18) !important;
          }
          .block-container .stButton > button[kind="secondary"],
          .block-container .stButton > button {
            background:
              linear-gradient(180deg, rgba(18, 44, 72, .96), rgba(11, 28, 48, .94)) !important;
            color: #f2fbff !important;
            border: 1px solid rgba(125,243,255,.26) !important;
            box-shadow: none !important;
            border-radius: 12px !important;
            min-height: 42px !important;
            padding: 0 16px !important;
          }
          .block-container .stButton > button:hover {
            border-color: rgba(125,243,255,.38) !important;
            background:
              linear-gradient(180deg, rgba(22, 53, 86, .98), rgba(12, 31, 54, .96)) !important;
          }
          .block-container [data-testid="stFileUploader"] {
            margin: 0 0 0;
            color: #f5fbff !important;
          }
          .block-container [data-testid="stFileUploader"] > div {
            background: transparent !important;
            box-shadow: none !important;
            border: 0 !important;
          }
          .block-container [data-testid="stFileUploader"] [data-testid="stFileUploaderDropzone"] + div {
            display: none !important;
          }
          .block-container [data-testid="stFileUploaderFileList"] {
            display: none !important;
            height: 0 !important;
            min-height: 0 !important;
            margin: 0 !important;
            padding: 0 !important;
            overflow: hidden !important;
          }
          .block-container [data-testid="stFileUploader"] label,
          .block-container [data-testid="stFileUploader"] label *,
          .block-container [data-testid="stFileUploader"] small,
          .block-container [data-testid="stFileUploader"] p,
          .block-container [data-testid="stFileUploader"] span {
            color: rgba(232, 246, 255, .86) !important;
          }
          .block-container [data-testid="stFileUploader"] section,
          .block-container [data-testid="stFileUploaderDropzone"] {
            background:
              linear-gradient(180deg, rgba(8, 20, 34, .90), rgba(5, 15, 28, .94)) !important;
            border: 1px dashed rgba(125,243,255,.24) !important;
            border-radius: 0 0 18px 18px !important;
            box-shadow:
              inset 0 1px 0 rgba(255,255,255,.06),
              inset 0 0 0 1px rgba(255,255,255,.02) !important;
            min-height: 152px !important;
          }
          .block-container [data-testid="stFileUploaderDropzone"] * {
            color: #edf8ff !important;
          }
          .block-container [data-testid="stFileUploaderDropzone"] {
            padding: 28px 24px !important;
          }
          .block-container [data-testid="stFileUploaderDropzone"] button,
          .block-container [data-testid="stFileUploader"] button {
            background: linear-gradient(180deg, rgba(20, 52, 84, .98), rgba(11, 31, 54, .96)) !important;
            border: 1px solid rgba(125,243,255,.26) !important;
            color: #eef9ff !important;
            box-shadow: none !important;
            border-radius: 999px !important;
            min-height: 42px !important;
            padding: 0 18px !important;
          }
          .block-container [data-testid="stFileUploader"] button [data-testid="stIconMaterial"] {
            width: 16px !important;
            overflow: hidden !important;
            color: transparent !important;
            font-size: 0 !important;
          }
          .block-container [data-testid="stFileUploader"] button [data-testid="stIconMaterial"]::after {
            content: "+";
            color: #eef9ff;
            font-size: 17px;
            line-height: 16px;
            font-weight: 760;
          }
          .block-container [data-testid="stFileUploader"] button [data-testid="stMarkdownContainer"] p {
            color: transparent !important;
            font-size: 0 !important;
          }
          .block-container [data-testid="stFileUploader"] button [data-testid="stMarkdownContainer"] p::after {
            content: "选择文件";
            color: #eef9ff;
            font-size: 13px;
            font-weight: 760;
            line-height: 1;
          }
          .block-container [data-testid="stFileUploaderDropzoneInstructions"] {
            color: transparent !important;
            font-size: 0 !important;
          }
          .block-container [data-testid="stFileUploaderDropzoneInstructions"]::after {
            content: "拖入电源设计文档，或点击按钮批量选择文件";
            color: #d6ebf6;
            font-size: 13px;
            line-height: 1.45;
            font-weight: 560;
          }
          .block-container [data-testid="stFileUploaderDropzoneInstructions"] * {
            color: transparent !important;
            font-size: 0 !important;
          }
          .block-container [data-testid="stFileUploaderDropzone"] button svg,
          .block-container [data-testid="stFileUploader"] button svg {
            color: #eef9ff !important;
            fill: #eef9ff !important;
          }
          .block-container [data-testid="stFileUploaderFile"] {
            display: none !important;
          }
          .block-container [data-testid="stFileUploader"] [data-testid="stFileUploaderDropzone"] {
            margin-bottom: 0 !important;
          }
          .block-container [data-testid="stFileUploader"] [data-testid="stFileUploaderDropzone"] > div:last-child {
            display: none !important;
          }
          .block-container [data-testid="stFileUploader"] [data-testid="stFileUploaderFileData"],
          .block-container [data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"] {
            display: none !important;
          }
          div[data-testid="stRadio"] > label {
            color: rgba(229, 244, 252, .92) !important;
            font-size: 12px !important;
            letter-spacing: .08em;
            text-transform: uppercase;
          }
          .block-container div[data-testid="stRadio"]:has([role="radiogroup"][aria-label="工作视图"]) [role="radiogroup"] {
            display: flex;
            flex-wrap: wrap;
            gap: 10px;
            margin-bottom: 18px;
          }
          .block-container div[data-testid="stRadio"]:has([role="radiogroup"][aria-label="工作视图"]) [role="radio"] {
            min-height: 42px;
            padding: 0 18px;
            border-radius: 999px;
            border: 1px solid rgba(125,243,255,.22);
            background: rgba(10, 24, 40, .74);
            display: inline-flex;
            align-items: center;
          }
          .block-container div[data-testid="stRadio"]:has([role="radiogroup"][aria-label="工作视图"]) [role="radio"][aria-checked="true"] {
            background:
              linear-gradient(180deg, rgba(18, 44, 72, .96), rgba(11, 28, 48, .94));
            border-color: rgba(125,243,255,.34);
            box-shadow: 0 14px 32px rgba(0, 0, 0, .18);
          }
          .block-container div[data-testid="stRadio"]:has([role="radiogroup"][aria-label="工作视图"]) [role="radio"][aria-checked="true"] label p {
            color: #f2fbff !important;
          }
          .block-container div[data-testid="stRadio"]:has([role="radiogroup"][aria-label="系统分区"]) [role="radiogroup"] {
            display: flex;
            flex-wrap: wrap;
            gap: 8px;
            margin: 2px 0 18px;
            padding: 6px;
            border-radius: 18px;
            background: rgba(8, 20, 34, .68);
            border: 1px solid rgba(125,243,255,.10);
          }
          .block-container div[data-testid="stRadio"]:has([role="radiogroup"][aria-label="系统分区"]) [role="radio"] {
            min-height: 38px;
            padding: 0 16px;
            border-radius: 999px;
            border: 1px solid transparent;
            background: transparent;
            display: inline-flex;
            align-items: center;
          }
          .block-container div[data-testid="stRadio"]:has([role="radiogroup"][aria-label="系统分区"]) [role="radio"][aria-checked="true"] {
            background:
              linear-gradient(180deg, rgba(18, 44, 72, .96), rgba(11, 28, 48, .94));
            border-color: rgba(125,243,255,.30);
          }
          .block-container div[data-testid="stRadio"]:has([role="radiogroup"][aria-label="系统分区"]) [role="radio"][aria-checked="true"] label p {
            color: #f2fbff !important;
          }
          .block-container div[data-testid="stRadio"] [role="radio"] label p {
            font-size: 14px !important;
            font-weight: 680 !important;
            color: rgba(234, 246, 252, .94) !important;
          }
          .stCaption {
            color: #e2f2fb !important;
            font-size: 15px !important;
          }
          .stCaption p {
            color: #e2f2fb !important;
          }
          div[data-testid="stStatusWidget"] {
            background: rgba(6, 17, 31, .76) !important;
            border: 1px solid rgba(125,243,255,.14) !important;
            border-radius: 14px !important;
          }
          div[data-testid="stStatusWidget"] * {
            color: #eef9ff !important;
          }
          .hero-band {
            position: relative;
            overflow: hidden;
            background:
              radial-gradient(circle at 14% 18%, rgba(46, 230, 255, .16), transparent 30%),
              radial-gradient(circle at 82% 22%, rgba(255, 184, 77, .10), transparent 26%),
              linear-gradient(135deg, rgba(4, 12, 25, .98), rgba(2, 7, 16, .94));
            background-size: cover;
            background-position: center;
            color: #ffffff;
            min-height: auto;
            padding: clamp(28px, 4vw, 56px);
            border-radius: 24px;
            margin: 0 0 28px;
            border: 1px solid rgba(226, 249, 255, .14);
            box-shadow:
              0 38px 110px rgba(0, 0, 0, .42),
              0 0 0 1px rgba(255,255,255,.04),
              inset 0 1px 0 rgba(255,255,255,.12);
            transform: translateY(0);
          }
          .hero-band::before {
            content: "";
            position: absolute;
            width: 58vw;
            height: 58vw;
            left: -26vw;
            bottom: -42vw;
            border: 1px solid rgba(125,243,255,.34);
            border-radius: 50%;
            box-shadow:
              0 0 24px rgba(46, 230, 255, .20),
              inset 0 0 42px rgba(46, 230, 255, .10);
            opacity: .34;
          }
          .hero-band::after {
            content: "";
            position: absolute;
            inset: 0;
            background:
              linear-gradient(90deg, rgba(125,243,255,.08) 1px, transparent 1px),
              linear-gradient(180deg, rgba(125,243,255,.06) 1px, transparent 1px);
            background-size: 52px 52px;
            opacity: .14;
            pointer-events: none;
          }
          .hero-content {
            position: relative;
            z-index: 1;
            display: grid;
            grid-template-columns: minmax(340px, 520px) minmax(580px, 1fr);
            gap: clamp(30px, 5vw, 72px);
            align-items: center;
            min-height: auto;
          }
          .hero-copy {
            align-self: center;
            max-width: 520px;
            padding: 0;
          }
          .hero-band-compact {
            padding: clamp(22px, 3vw, 34px);
            margin: 0 0 22px;
          }
          .hero-content-compact {
            grid-template-columns: minmax(340px, 1fr) minmax(340px, 520px);
            gap: clamp(24px, 4vw, 44px);
          }
          .hero-status-compact {
            grid-template-columns: repeat(3, 1fr);
            margin-top: 22px;
          }
          .video-stage-compact {
            max-height: 320px;
          }
          .system-kicker {
            display: inline-flex;
            color: rgba(212, 240, 251, .9);
            font-size: 12px;
            font-weight: 720;
            letter-spacing: .10em;
            text-transform: uppercase;
            margin-bottom: 16px;
          }
          .hero-band h1 {
            max-width: 560px;
            margin: 0 0 18px;
            font-family: -apple-system, BlinkMacSystemFont, "SF Pro Display", "Avenir Next", "Segoe UI", sans-serif;
            font-size: clamp(48px, 5.6vw, 86px);
            line-height: .94;
            font-weight: 780;
            letter-spacing: -.055em;
            color: #f2fbff;
            text-shadow: 0 14px 42px rgba(46, 230, 255, .14);
          }
          .hero-band p {
            max-width: 500px;
            margin: 0;
            color: rgba(237, 244, 249, .92);
            line-height: 1.68;
            font-size: 16px;
          }
          .mission-chips {
            display: flex;
            flex-wrap: wrap;
            gap: 8px;
            margin-top: 20px;
          }
          .mission-chips span {
            display: inline-flex;
            min-height: 32px;
            align-items: center;
            padding: 6px 11px;
            border-radius: 999px;
            border: 1px solid rgba(125, 243, 255, .26);
            background: rgba(226, 249, 255, .06);
            color: rgba(244, 249, 252, .96);
            font-size: 12px;
            font-weight: 760;
          }
          .hero-status {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 10px;
            margin-top: 26px;
          }
          .hero-status div {
            display: grid;
            gap: 7px;
            padding: 14px;
            border: 1px solid rgba(226,249,255,.14);
            border-radius: 14px;
            background: rgba(226,249,255,.055);
            font-size: 13px;
          }
          .hero-status div:last-child {
            border-bottom: 1px solid rgba(226,249,255,.14);
          }
          .hero-status span {
            color: rgba(212, 232, 241, .86);
          }
          .hero-status strong {
            color: #ffffff;
            font-weight: 700;
          }
          .video-stage {
            position: relative;
            aspect-ratio: 16 / 9;
            width: 100%;
            max-height: 560px;
            border-radius: 24px;
            border: 1px solid rgba(226, 249, 255, .16);
            background: rgba(4, 14, 29, .46);
            box-shadow:
              0 34px 90px rgba(0,0,0,.34),
              0 0 48px rgba(46,230,255,.08),
              inset 0 1px 0 rgba(255,255,255,.12);
            overflow: hidden;
            transform: none;
          }
          .video-stage::before {
            content: "";
            position: absolute;
            inset: 10px;
            border: 1px solid rgba(255, 255, 255, .10);
            border-radius: 18px;
            pointer-events: none;
            z-index: 2;
          }
          .video-stage::after {
            content: "";
            position: absolute;
            inset: 0;
            z-index: 3;
            pointer-events: none;
            box-shadow: inset 0 -70px 90px rgba(0,0,0,.16);
          }
          .hero-video {
            position: absolute;
            inset: 0;
            width: 100%;
            height: 100%;
            object-fit: contain;
            object-position: center;
            z-index: 1;
            opacity: 1;
            filter: brightness(1.18) contrast(1.04) saturate(1.03);
          }
          .video-glass-line {
            position: absolute;
            inset: 0;
            z-index: 2;
            pointer-events: none;
            background:
              linear-gradient(90deg, rgba(125,243,255,.09) 1px, transparent 1px),
              linear-gradient(180deg, rgba(125,243,255,.07) 1px, transparent 1px);
            background-size: 54px 54px;
            opacity: .08;
          }
          .metric-tile {
            min-height: 94px;
            border: 1px solid var(--pd-line);
            background: var(--pd-surface);
            border-radius: var(--pd-radius);
            padding: 16px 17px;
            display: flex;
            flex-direction: column;
            justify-content: space-between;
            box-shadow: 0 18px 42px rgba(0, 0, 0, .20), inset 0 1px 0 rgba(255,255,255,.08);
          }
          .metric-tile span {
            color: #b8d1df;
            font-size: 13px;
            font-weight: 650;
          }
          .metric-tile strong {
            color: var(--pd-ink);
            font-size: 24px;
            line-height: 1.1;
          }
          .kb-card {
            border: 1px solid var(--pd-line);
            background: var(--pd-surface);
            border-radius: var(--pd-radius);
            padding: 15px 16px;
            margin-bottom: 10px;
            box-shadow: 0 16px 38px rgba(0, 0, 0, .18), inset 0 1px 0 rgba(255,255,255,.07);
          }
          .kb-card-title {
            font-weight: 700;
            color: var(--pd-ink);
            margin-bottom: 7px;
          }
          .kb-card-body {
            color: #e0edf5;
            line-height: 1.65;
            font-size: 14px;
          }
          .kb-card-meta {
            margin-top: 10px;
            color: #9bb4c2;
            font-size: 12px;
          }
          .section-label {
            display: flex;
            justify-content: space-between;
            align-items: end;
            gap: 20px;
            margin: 18px 0 12px;
          }
          .section-label h2 {
            margin: 0;
            font-size: 20px;
            line-height: 1.15;
            color: var(--pd-ink);
          }
          .section-label p {
            max-width: 58ch;
            margin: 0;
            color: #c4d8e4;
            font-size: 13px;
            line-height: 1.55;
          }
          .source-inline {
            margin-top: 12px;
            padding-top: 8px;
            border-top: 1px solid rgba(125,243,255,.14);
          }
          .source-inline span {
            display: inline-flex;
            color: var(--pd-accent-dark);
            font-size: 12px;
            font-weight: 700;
          }
          .summary-strip {
            margin: 4px 0 20px;
          }
          .summary-chip {
            min-height: 88px;
            border-radius: 16px;
            padding: 14px 16px;
            border: 1px solid rgba(125,243,255,.12);
            background: rgba(8, 20, 34, .72);
          }
          .summary-chip span {
            display: block;
            font-size: 12px;
            color: #b6cedd;
            margin-bottom: 10px;
          }
          .summary-chip strong {
            display: block;
            font-size: 24px;
            color: #f5fbff;
            line-height: 1.08;
          }
          .tag-row {
            display: flex;
            flex-wrap: wrap;
            gap: 8px;
          }
          .tag {
            display: inline-flex;
            align-items: center;
            min-height: 28px;
            padding: 4px 10px;
            border-radius: 999px;
            background: rgba(46, 230, 255, .10);
            color: #d9f8ff;
            font-size: 13px;
            border: 1px solid rgba(125,243,255,.26);
          }
          .json-box pre {
            max-height: 680px;
          }
          .flow-lane-wrap {
            margin-bottom: 18px;
          }
          .flow-step {
            position: relative;
            min-height: 148px;
            border: 1px solid var(--pd-line);
            background: var(--pd-surface);
            border-radius: var(--pd-radius);
            padding: 14px;
            overflow: hidden;
          }
          .flow-step::before {
            content: "";
            position: absolute;
            inset: 0;
            background: linear-gradient(110deg, transparent 0%, rgba(46, 230, 255, .18) 46%, transparent 78%);
            transform: translateX(-100%);
            animation: pd-scan 2.8s ease-in-out infinite;
          }
          .flow-step span,
          .flow-step strong,
          .flow-step small {
            position: relative;
            z-index: 1;
          }
          .flow-step span {
            display: inline-flex;
            color: var(--pd-accent-dark);
            font-size: 12px;
            font-weight: 800;
            margin-bottom: 22px;
          }
          .flow-step strong {
            display: block;
            color: var(--pd-ink);
            font-size: 15px;
            line-height: 1.25;
            margin-bottom: 8px;
          }
          .flow-step small {
            color: var(--pd-muted);
            font-size: 12px;
          }
          .table-shell {
            width: 100%;
            overflow-x: auto;
            border: 1px solid rgba(226,249,255,.14);
            border-radius: 14px;
            background: rgba(5, 14, 28, .88);
            box-shadow: 0 20px 50px rgba(0,0,0,.22), inset 0 1px 0 rgba(255,255,255,.06);
          }
          .dark-table {
            width: 100%;
            min-width: 760px;
            border-collapse: collapse;
            color: #dff8ff;
            font-size: 13px;
          }
          .dark-table th {
            position: sticky;
            top: 0;
            background: rgba(12, 30, 52, .96);
            color: rgba(198, 240, 255, .72);
            text-align: left;
            font-weight: 760;
            padding: 12px 14px;
            border-bottom: 1px solid rgba(226,249,255,.12);
          }
          .dark-table td {
            padding: 12px 14px;
            border-bottom: 1px solid rgba(226,249,255,.08);
            color: rgba(226, 249, 255, .86);
            white-space: nowrap;
          }
          .dark-table tr:last-child td {
            border-bottom: 0;
          }
          .code-panel {
            overflow: auto;
            margin: 0;
            padding: 18px 20px;
            border: 1px solid rgba(226,249,255,.14);
            border-radius: 14px;
            background: rgba(5, 14, 28, .92);
            color: #dff8ff;
            box-shadow: 0 20px 50px rgba(0,0,0,.22), inset 0 1px 0 rgba(255,255,255,.06);
          }
          .code-panel pre {
            margin: 0;
            background: transparent !important;
          }
          .code-panel code {
            color: #dff8ff;
            font-family: "SF Mono", "Menlo", "Consolas", monospace;
            font-size: 13px;
            line-height: 1.62;
            white-space: pre;
          }
          .page-heatmap {
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(58px, 1fr));
            gap: 7px;
            margin-bottom: 18px;
          }
          .page-cell {
            min-height: 58px;
            border: 1px solid rgba(125, 243, 255, calc(.16 + var(--level) * .42));
            background:
              linear-gradient(180deg, rgba(46, 230, 255, calc(.08 + var(--level) * .44)), rgba(255, 184, 77, calc(.04 + var(--level) * .18))),
              var(--pd-surface);
            border-radius: var(--pd-radius);
            padding: 8px;
            display: flex;
            flex-direction: column;
            justify-content: space-between;
          }
          .page-cell span {
            color: var(--pd-muted);
            font-size: 11px;
            font-weight: 700;
          }
          .page-cell strong {
            color: var(--pd-ink);
            font-size: 18px;
            line-height: 1;
          }
          .bar-panel {
            border: 1px solid var(--pd-line);
            background: var(--pd-surface);
            border-radius: var(--pd-radius);
            padding: 14px;
            min-height: 268px;
          }
          .signal-bar {
            margin-bottom: 13px;
          }
          .signal-bar div {
            display: flex;
            justify-content: space-between;
            gap: 16px;
            color: var(--pd-muted);
            font-size: 13px;
            margin-bottom: 7px;
          }
          .signal-bar strong {
            color: var(--pd-ink);
          }
          .signal-bar i {
            display: block;
            height: 10px;
            border-radius: 999px;
            background: linear-gradient(90deg, var(--pd-accent), var(--pd-green));
            box-shadow: inset 0 0 0 1px rgba(255,255,255,.25);
            animation: pd-grow .75s ease-out both;
          }
          .node-map {
            position: relative;
            min-height: 268px;
            border: 1px solid var(--pd-line);
            background:
              linear-gradient(90deg, rgba(125, 243, 255, .08) 1px, transparent 1px),
              linear-gradient(180deg, rgba(125, 243, 255, .06) 1px, transparent 1px),
              var(--pd-surface);
            background-size: 34px 34px;
            border-radius: var(--pd-radius);
            padding: 14px;
            display: grid;
            grid-template-columns: repeat(3, 1fr);
            gap: 12px;
            align-items: center;
          }
          .node-map::before,
          .node-map::after {
            content: "";
            position: absolute;
            left: 11%;
            right: 11%;
            top: 50%;
            height: 1px;
            background: rgba(46, 230, 255, .28);
          }
          .node-map::after {
            top: 26%;
            transform: rotate(18deg);
            transform-origin: center;
          }
          .knowledge-node {
            position: relative;
            z-index: 1;
            min-height: 78px;
            border: 1px solid rgba(125, 243, 255, .30);
            background: rgba(6, 20, 38, .92);
            border-radius: var(--pd-radius);
            padding: 12px;
            transform: scale(var(--scale));
            transform-origin: center;
            box-shadow: 0 16px 34px rgba(0, 0, 0, .24), 0 0 22px rgba(46, 230, 255, .08);
          }
          .knowledge-node strong {
            display: block;
            color: var(--pd-accent-dark);
            font-size: 22px;
            line-height: 1;
            margin-bottom: 8px;
          }
          .knowledge-node span {
            color: var(--pd-ink);
            font-size: 12px;
            font-weight: 700;
          }
          .empty-note {
            color: var(--pd-muted);
            font-size: 13px;
            padding: 10px;
          }
          @keyframes pd-scan {
            0% { transform: translateX(-100%); }
            48%, 100% { transform: translateX(120%); }
          }
          @keyframes pd-grow {
            from { transform: scaleX(.18); transform-origin: left; }
            to { transform: scaleX(1); transform-origin: left; }
          }
          @keyframes star-drift {
            from { transform: translate3d(0, 0, 0); }
            to { transform: translate3d(-96px, 96px, 0); }
          }
          @keyframes orbital-scan {
            0%, 100% { background-position: 0 0, 0 0, -80% 0; }
            50% { background-position: 0 0, 0 0, 120% 0; }
          }
          @keyframes orbit-pulse {
            0%, 100% { opacity: .42; transform: scale(.98); }
            50% { opacity: .78; transform: scale(1.02); }
          }
          @media (prefers-reduced-motion: reduce) {
            .flow-step::before,
            .signal-bar i,
            .stApp::before,
            .hero-band::before,
            .hero-band::after {
              animation: none;
            }
          }
          .empty-workbench {
            display: grid;
            grid-template-columns: minmax(0, 1fr) minmax(280px, 420px);
            gap: 22px;
            align-items: stretch;
            border: 1px solid var(--pd-line);
            background:
              linear-gradient(135deg, rgba(7, 18, 34, .88), rgba(5, 12, 24, .78)),
              radial-gradient(circle at 100% 0%, rgba(46, 230, 255, .12), transparent 32%);
            border-radius: var(--pd-radius);
            padding: 22px;
            margin-top: -10px;
            box-shadow: 0 24px 70px rgba(0,0,0,.26), inset 0 1px 0 rgba(255,255,255,.08);
          }
          .empty-workbench .system-kicker {
            color: var(--pd-accent-dark);
            margin-bottom: 10px;
          }
          .empty-workbench h2 {
            margin: 0 0 10px;
            font-size: clamp(26px, 4vw, 42px);
            line-height: 1.02;
          }
          .empty-workbench p {
            margin: 0;
            max-width: 58ch;
            color: var(--pd-muted);
            line-height: 1.65;
          }
          .empty-grid {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 8px;
          }
          .empty-grid div {
            border: 1px solid var(--pd-line);
            background: var(--pd-surface-muted);
            border-radius: var(--pd-radius);
            padding: 13px;
            min-height: 104px;
          }
          .empty-grid strong {
            display: block;
            color: var(--pd-accent-dark);
            font-size: 24px;
            margin-bottom: 12px;
          }
          .empty-grid span {
            color: var(--pd-ink);
            font-size: 13px;
            line-height: 1.35;
          }
          div[data-testid="stAlert"] {
            border-radius: var(--pd-radius);
            border-color: var(--pd-line);
            background: rgba(7, 18, 34, .82);
          }
          div[data-testid="stTabs"] button {
            border-radius: var(--pd-radius) var(--pd-radius) 0 0;
          }
          .stDownloadButton button,
          .stButton button {
            border-radius: var(--pd-radius);
            min-height: 40px;
            font-weight: 700;
          }
          div[data-testid="stCodeBlock"] pre {
            background: rgba(5, 14, 28, .92) !important;
            border: 1px solid rgba(226,249,255,.14);
            border-radius: 14px;
          }
          div[data-testid="stCodeBlock"] code {
            color: #dff8ff !important;
          }
          @media (max-width: 860px) {
            .hero-content,
            .empty-workbench {
              grid-template-columns: 1fr;
            }
            .hero-band {
              min-height: auto;
              transform: none;
            }
            .hero-content {
              min-height: auto;
            }
            .hero-status {
              grid-column: 1;
            }
            .hero-content-compact {
              grid-template-columns: 1fr;
            }
            .hero-status-compact {
              grid-template-columns: 1fr;
            }
            .result-header-grid {
              grid-template-columns: 1fr 1fr;
            }
            .hero-band h1 {
              font-size: clamp(44px, 12vw, 72px);
            }
            .hero-band p {
              font-size: 15px;
            }
            .mission-chips {
              margin-top: 16px;
            }
            .upload-surface-head {
              display: grid;
              gap: 10px;
            }
            .upload-surface-badge {
              width: fit-content;
            }
            .upload-surface-metrics {
              grid-template-columns: 1fr;
            }
            .upload-surface-foot {
              padding-top: 10px;
            }
            .video-stage-compact {
              max-height: none;
            }
            .empty-workbench {
              margin-top: -8px;
            }
            .flow-lane,
            .node-map {
              grid-template-columns: 1fr;
            }
            .section-label {
              display: block;
            }
            .section-label p {
              margin-top: 6px;
            }
          }
        </style>
        """
    st.markdown(
        css.replace("__ORBIT_LOOP__", asset_data_url(ORBIT_LOOP_PATH, "image/webp")),
        unsafe_allow_html=True,
    )


def render_overview(payload: dict[str, Any]) -> None:
    summary = payload["summary"]
    power_system = payload.get("power_system", {})
    cols = st.columns(4)
    with cols[0]:
        render_metric("系统模型", power_system.get("schema", "未生成"))
    with cols[1]:
        render_metric("结构化指标", len(power_system.get("metrics", [])))
    with cols[2]:
        render_metric("设计约束", len(power_system.get("constraints", [])))
    with cols[3]:
        render_metric("测试矩阵", len(power_system.get("test_cases", [])))

    topology = power_system.get("topology_scheme", {})
    section_label("系统拓扑方案", "按航天电源系统语义抽取的架构，不再只是关键词命中。")
    context_pages = topology.get("context_pages", [])
    topology_meta = f"置信度：{topology.get('confidence', 0)}"
    if context_pages:
        topology_meta += f" · 主要页码：{' / '.join(f'P{page}' for page in context_pages[:4])}"
    card("系统架构", f"{escape(topology.get('architecture', '未识别'))}<br>{escape(topology_meta)}")
    feature_rows = [
        [item["id"], item["label"], item.get("summary", ""), item["source"].get("page")]
        for item in topology.get("features", [])
    ]
    if feature_rows:
        render_dark_table(["特征", "名称", "归纳说明", "页码"], feature_rows)

    section_label("系统分区", "仿照样例把原文组织成任务、轨道、电气指标、拓扑、太阳翼、蓄电池、控制设备、可靠性等模块。")
    summary_rows = [
        [item["label"], item["items"], " / ".join(item["preview"])]
        for item in power_system.get("section_summary", [])
    ]
    if summary_rows:
        render_dark_table(["section", "items", "preview"], summary_rows)

    section_label("关键词", "从文档正文中提取的高频技术词，用来快速判断文档主题。")
    tags = "".join(f'<span class="tag">{escape(keyword)}</span>' for keyword in summary["keyword_cloud"])
    st.markdown(f'<div class="tag-row">{tags}</div>', unsafe_allow_html=True)

    section_label("识别到的拓扑", "按出现次数排序，并保留可追溯的原文片段。")
    if not payload["topologies"]:
        st.info("暂未识别到明确拓扑，可以上传 datasheet、application note 或设计指南类文档再试。")
    for item in payload["topologies"]:
        card(
            item["label"],
            f"出现次数：{item['mentions']}<br>置信度：{item['confidence']}",
            render_source(item["source"]),
        )


def render_system_model(power_system: dict[str, Any]) -> None:
    section_label("航天电源系统模型", "这是面向知识库入库的主结果：系统级结构、指标、约束和交付物。")
    top_cols = st.columns(4)
    with top_cols[0]:
        render_metric("模型版本", power_system.get("schema", "n/a"))
    with top_cols[1]:
        render_metric("功能项", len(power_system.get("functional_requirements", [])))
    with top_cols[2]:
        render_metric("指标数量", len(power_system.get("metrics", [])))
    with top_cols[3]:
        render_metric("测试项数量", len(power_system.get("test_cases", [])))

    topology = power_system.get("topology_scheme", {})
    section_label("拓扑方案", "系统级供电架构、调节方式和关键功能链路。")
    topology_pages = topology.get("context_pages", [])
    topology_hint = ""
    if topology_pages:
        topology_hint = f'<div class="result-caption">主要依据页：{" / ".join(f"P{page}" for page in topology_pages[:4])}</div>'
    card("拓扑摘要", escape(topology.get("architecture", "未识别")) + topology_hint)
    for feature in topology.get("features", []):
        body = escape(feature.get("summary", feature["label"]))
        card(feature["label"], body, render_source(feature["source"]))

    functional_requirements = power_system.get("functional_requirements", [])
    if functional_requirements:
        section_label("功能定义", "按样例中的任务和功能定义模块整理功能项与关键配置要求。")
        render_dark_table(
            ["功能项", "功能描述", "关键配置/要求", "页码"],
            [
                [item["name"], item["description"], item["configuration"], item["source"].get("page")]
                for item in functional_requirements
            ],
        )

    section_label("结构化指标", "从表格和段落中归并出的设计指标，保留原文依据。")
    metric_rows = [
        [
            SYSTEM_SECTION_LABELS.get(item["section"], item["section"]),
            item["name"],
            item["value"],
            item.get("note", ""),
            ", ".join(item["units"]),
            item["source"].get("page"),
        ]
        for item in power_system.get("metrics", [])
    ]
    if metric_rows:
        render_dark_table(["分区", "名称", "指标值", "备注/判据", "单位", "页码"], metric_rows)
    else:
        st.info("没有识别到结构化指标。")

    section_label("设计约束", "把指标中带有范围、上下限、稳定性、寿命、恢复时间等语义的条目提炼为约束。")
    for item in power_system.get("constraints", [])[:50]:
        body = f"<strong>{escape(item['requirement'])}</strong><br>{escape(item['rationale'])}"
        card(f"{SYSTEM_SECTION_LABELS.get(item['category'], item['category'])} / {item['name']}", body, render_source(item["source"]))

    deliverables = power_system.get("deliverables", [])
    if deliverables:
        section_label("研试文件与交付物", "适合直接进入知识库的文档清单。")
        render_dark_table(
            ["类别", "名称", "文件类型", "页码"],
            [[item["category"], item["name"], item["type"], item["source"].get("page")] for item in deliverables],
        )

    test_cases = power_system.get("test_cases", [])
    if test_cases:
        section_label("测试验证矩阵", "从 PWR-TC / EQP-TC 等验证表中抽取测试项、前置条件和合格判据。")
        render_dark_table(
            ["编号", "测试项", "前置条件", "考核对象", "合格判据", "页码"],
            [
                [
                    item["id"],
                    item["name"],
                    item["precondition"][:64],
                    item["target_rule"][:48],
                    item["acceptance_criteria"][:72],
                    item["source"].get("page"),
                ]
                for item in test_cases[:24]
            ],
        )


def render_section_browser(power_system: dict[str, Any]) -> None:
    sections = power_system.get("sections", {})
    if not sections:
        st.info("没有可浏览的系统分区。")
        return
    labels = {SYSTEM_SECTION_LABELS.get(section, section): section for section in sections}
    selected_label = render_section_selector(list(labels), list(labels)[0])
    section = labels[selected_label]
    for item in sections.get(section, []):
        title = item.get("title") or selected_label
        body = escape(item.get("content") or "")
        card(title, body, render_source(item["source"]))


def render_parameters(parameters: list[dict[str, Any]]) -> None:
    if not parameters:
        st.info("没有抽取到标准参数。")
        return
    section_label("参数表", "抽取 Vin、Vout、Iout、fsw、效率、纹波等工程参数。")
    render_dark_table(
        ["name", "raw", "value", "unit", "page", "confidence"],
        [
            [
                item["name"],
                item["raw_label"],
                item["value"],
                item["unit"],
                item["source"].get("page"),
                item["confidence"],
            ]
            for item in parameters
        ],
    )
    section_label("参数溯源", "前 12 条参数以卡片形式展示，方便快速回看上下文。")
    for item in parameters[:12]:
        body = f"<strong>{escape(item['value'])} {escape(item['unit'])}</strong><br>原始标签：{escape(item['raw_label'])}"
        card(item["name"], body, render_source(item["source"]))


def render_formulas(formulas: list[dict[str, Any]]) -> None:
    if not formulas:
        st.info("没有抽取到公式。")
        return
    section_label("公式", "保留表达式、变量列表和原文位置，后续可以接入公式校验。")
    for item in formulas:
        variables = ", ".join(item["variables"]) if item["variables"] else "未识别"
        body = f"<code>{escape(item['expression'])}</code><br>变量：{escape(variables)}<br>置信度：{item['confidence']}"
        card(item["name"], body, render_source(item["source"]))


def render_components(components: list[dict[str, Any]]) -> None:
    if not components:
        st.info("没有抽取到器件信息。")
        return
    section_label("器件", "统计控制器、MOSFET、电感、电容、驱动器、吸收电路等器件线索。")
    for item in components:
        notes = "<br>".join(escape(note) for note in item["selection_notes"]) or "暂无明确选型句子"
        body = f"出现次数：{item['mentions']}<br>选型线索：<br>{notes}"
        card(item["type"], body, render_source(item["source"]))


def render_rules(rules: list[dict[str, Any]], category: str | None = None) -> None:
    filtered = [rule for rule in rules if category is None or rule["category"] == category]
    if not filtered:
        st.info("这个分类下暂时没有抽取结果。")
        return
    section_label("设计规则", "规则按 layout、EMI、thermal、protection、selection 和 troubleshooting 分类。")
    for rule in filtered:
        body = escape(rule["rule"])
        meta = f"category={rule['category']} / severity={rule['severity']} / confidence={rule['confidence']}"
        card(rule["category"], body, meta + render_source(rule["source"]))


def main() -> None:
    st.set_page_config(page_title=APP_TITLE, page_icon="PD", layout="wide")
    inject_css()
    st.session_state.setdefault("powerdoc_parsed_items", [])
    st.session_state.setdefault("powerdoc_show_uploader", True)
    st.session_state.setdefault("powerdoc_active_doc_index", 0)
    st.session_state.setdefault("powerdoc_active_view", "System Model")
    st.session_state.setdefault("powerdoc_active_section", "")
    st.session_state.setdefault("powerdoc_uploader_version", 0)
    hero_video_src = asset_data_url(ORBIT_VIDEO_PATH, "video/mp4")

    with st.sidebar:
        st.markdown(
            """
            <div class="sidebar-brand-shell">
              <strong>PowerDoc-KB</strong>
              <span>电源设计文档结构化工具</span>
            </div>
            """,
            unsafe_allow_html=True,
        )

    parsed_items = st.session_state.get("powerdoc_parsed_items", [])
    render_workspace_masthead(hero_video_src, bool(parsed_items), len(parsed_items))
    uploaded_files = render_upload_panel(
        st.session_state.get("powerdoc_show_uploader", True),
        len(parsed_items),
    )

    if uploaded_files:
        try:
            progress_lines = [f"准备处理 {len(uploaded_files)} 份文档"]
            render_processing_banner(progress_lines)
            parsed_items = []
            for idx, uploaded_file in enumerate(uploaded_files, start=1):
                progress_lines = [
                    f"读取文件结构 {idx}/{len(uploaded_files)}：{uploaded_file.name}",
                    "切分页码和文本段落",
                    "抽取系统模型、指标和约束",
                    "绑定原文页码并生成 JSON",
                ]
                render_processing_banner(progress_lines)
                doc = parse_upload(uploaded_file)
                payload = extract_knowledge(doc)
                parsed_items.append({"doc": doc, "payload": payload})
            st.session_state["powerdoc_parsed_items"] = parsed_items
            st.session_state["powerdoc_show_uploader"] = False
            st.session_state["powerdoc_active_doc_index"] = 0
            st.session_state["powerdoc_active_view"] = "System Model"
            st.session_state["powerdoc_active_section"] = ""
            st.rerun()
        except Exception as exc:
            st.error(f"解析失败：{exc}")
            st.stop()
    else:
        parsed_items = st.session_state.get("powerdoc_parsed_items", [])

    if not parsed_items:
        with st.sidebar:
            st.header("文档队列")
            render_sidebar_empty_state()
        render_empty_state()
        demo = {
            "schema": "aerospacepower_system.v1",
            "output": ["topology_scheme", "metrics", "constraints", "deliverables", "source_trace"],
            "extractor": "schema_driven_rule_v2",
        }
        section_label("JSON 输出预览")
        code_panel(json.dumps(demo, ensure_ascii=False, indent=2), max_height=360)
        return

    doc_cards = [
        {
            "name": item["payload"]["document"]["filename"],
            "pages": item["payload"]["document"]["pages"],
            "chars": item["payload"]["summary"]["character_count"],
        }
        for item in parsed_items
    ]
    default_doc_index = st.session_state.get("powerdoc_active_doc_index", 0)
    if default_doc_index >= len(doc_cards):
        default_doc_index = 0
    active_doc_index = render_document_switcher(doc_cards, default_doc_index)
    st.session_state["powerdoc_active_doc_index"] = active_doc_index
    active_doc = parsed_items[active_doc_index]["doc"]
    active_payload = parsed_items[active_doc_index]["payload"]
    with st.sidebar:
        st.header("当前文档")
        render_sidebar_active_document(doc_cards[active_doc_index])
        if len(doc_cards) > 1:
            st.header("已载入文档")
            render_sidebar_file_list(doc_cards, active_doc_index)

    render_result_header(doc_cards[active_doc_index], len(doc_cards), active_payload)

    enabled_views = ["System Model", "Section Browser", "Visual Parse", "Overview", "JSON"]
    active_view = render_view_selector(enabled_views)

    if active_view == "System Model":
        render_system_model(active_payload["power_system"])
    elif active_view == "Section Browser":
        render_section_browser(active_payload["power_system"])
    elif active_view == "Visual Parse":
        render_parse_visualization(active_payload, active_doc)
    elif active_view == "Overview":
        render_overview(active_payload)
    elif active_view == "Parameters":
        render_parameters(active_payload["parameters"])
    elif active_view == "Formulas":
        render_formulas(active_payload["formulas"])
    elif active_view == "Components":
        render_components(active_payload["components"])
    elif active_view == "Design Rules":
        categories = ["all"] + sorted({rule["category"] for rule in active_payload["design_rules"]})
        selected = st.selectbox("规则分类", categories, index=0)
        render_rules(active_payload["design_rules"], None if selected == "all" else selected)
    elif active_view == "JSON":
        json_text = json.dumps(active_payload, ensure_ascii=False, indent=2)
        st.download_button(
            "下载当前文档 JSON",
            data=json_text,
            file_name=f"{Path(active_payload['document']['filename']).stem}.powerdoc.json",
            mime="application/json",
            type="primary",
        )
        code_panel(json_text, max_height=720)


if __name__ == "__main__":
    main()
